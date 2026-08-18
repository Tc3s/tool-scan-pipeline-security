#!/usr/bin/env python3
"""
Advanced DOCX Template Analyzer Tool
Deep inspection of Microsoft Word (.docx) templates: page layout, margins,
headers/footers, watermarks, logo placements, table grid micro-architecture,
typography nuances, pagination controls, numbered lists, and field codes.
"""

import os
import sys
import json
import docx
from docx.shared import Inches, Pt, Twips, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# XML Namespace Mapping
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'v': 'urn:schemas-microsoft-com:vml',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
}

def dxa_to_pt(dxa):
    if dxa is None:
        return None
    try:
        return round(float(dxa) / 20.0, 2)
    except (ValueError, TypeError):
        return None

def emu_to_pt(emu):
    if emu is None:
        return None
    try:
        return round(float(emu) / 12700.0, 2)
    except (ValueError, TypeError):
        return None

def get_xml_attr(elem, xpath, attr_name):
    nodes = elem.xpath(xpath)
    if nodes:
        node = nodes[0]
        if hasattr(node, 'get'):
            return node.get(qn(attr_name)) or node.get(attr_name)
        return str(node)
    return None

def analyze_sections(doc):
    """Area 1: Document margins, orientation, and page sizes (w:pgSz, w:pgMar)"""
    sections_data = []
    for idx, sec in enumerate(doc.sections):
        sp = sec._sectPr
        
        pg_sz = sp.xpath('./w:pgSz')
        w_val = pg_sz[0].get(qn('w:w')) if pg_sz else None
        h_val = pg_sz[0].get(qn('w:h')) if pg_sz else None
        orient = pg_sz[0].get(qn('w:orient')) if pg_sz else 'portrait'

        top = dxa_to_pt(sec.top_margin.twips) if sec.top_margin else None
        bottom = dxa_to_pt(sec.bottom_margin.twips) if sec.bottom_margin else None
        left = dxa_to_pt(sec.left_margin.twips) if sec.left_margin else None
        right = dxa_to_pt(sec.right_margin.twips) if sec.right_margin else None
        header_dist = dxa_to_pt(sec.header_distance.twips) if sec.header_distance else None
        footer_dist = dxa_to_pt(sec.footer_distance.twips) if sec.footer_distance else None
        gutter = dxa_to_pt(sec.gutter.twips) if sec.gutter else 0.0

        cols_node = sp.xpath('./w:cols')
        cols_count = int(cols_node[0].get(qn('w:num')) or 1) if cols_node else 1
        cols_space = dxa_to_pt(cols_node[0].get(qn('w:space'))) if cols_node else None

        sec_info = {
            "section_index": idx + 1,
            "start_type": str(sec.start_type),
            "orientation": orient,
            "page_width_pt": dxa_to_pt(w_val),
            "page_height_pt": dxa_to_pt(h_val),
            "margins_pt": {
                "top": top,
                "bottom": bottom,
                "left": left,
                "right": right,
                "header": header_dist,
                "footer": footer_dist,
                "gutter": gutter
            },
            "columns": {
                "count": cols_count,
                "spacing_pt": cols_space
            },
            "different_first_page_header_footer": sec.different_first_page_header_footer
        }
        sections_data.append(sec_info)
    return sections_data

def inspect_header_footer_xml(hf_obj, hf_type):
    """Area 2: Header/Footer XML inspection, watermarks, logo placements"""
    if not hf_obj or hf_obj.is_linked_to_previous:
        return {
            "type": hf_type,
            "status": "linked_to_previous" if (hf_obj and hf_obj.is_linked_to_previous) else "empty"
        }

    elem = hf_obj._element
    paragraphs_info = []
    for p in hf_obj.paragraphs:
        if p.text.strip():
            paragraphs_info.append({
                "style": p.style.name if p.style else None,
                "alignment": str(p.alignment) if p.alignment else None,
                "text": p.text.strip()
            })

    drawings = []
    for d in elem.xpath('.//*[local-name()="drawing" or local-name()="shape" or local-name()="pict"]'):
        blips = d.xpath('.//*[local-name()="blip"]')
        r_id = blips[0].get(qn('r:embed')) if blips else None
        
        exts = d.xpath('.//*[local-name()="extent"]')
        cx = emu_to_pt(exts[0].get('cx')) if exts else None
        cy = emu_to_pt(exts[0].get('cy')) if exts else None

        drawings.append({
            "tag": d.tag.split('}')[-1],
            "embed_rid": r_id,
            "width_pt": cx,
            "height_pt": cy
        })

    watermarks = []
    v_shapes = elem.xpath('.//v:shape | .//*[local-name()="shape"]')
    for v_shape in v_shapes:
        style_attr = v_shape.get('style') or ''
        textpath = v_shape.xpath('.//v:textpath | .//*[local-name()="textpath"]')
        if textpath or 'watermark' in style_attr.lower() or 'powerpluswatermark' in (v_shape.get('id') or '').lower():
            wm_text = textpath[0].get('string') if textpath else None
            watermarks.append({
                "shape_id": v_shape.get('id'),
                "text": wm_text,
                "style": style_attr
            })

    fields = []
    for instr in elem.xpath('.//*[local-name()="instrText"]'):
        txt = (instr.text or "").strip()
        if txt:
            fields.append(txt)

    return {
        "type": hf_type,
        "is_linked_to_previous": hf_obj.is_linked_to_previous,
        "paragraphs": paragraphs_info,
        "drawings_count": len(drawings),
        "drawings": drawings,
        "watermarks": watermarks,
        "field_codes": fields
    }

def analyze_headers_footers(doc):
    hf_data = []
    for idx, sec in enumerate(doc.sections):
        sec_hf = {
            "section_index": idx + 1,
            "header_default": inspect_header_footer_xml(sec.header, "header_default"),
            "header_first": inspect_header_footer_xml(sec.first_page_header, "header_first"),
            "header_even": inspect_header_footer_xml(sec.even_page_header, "header_even"),
            "footer_default": inspect_header_footer_xml(sec.footer, "footer_default"),
            "footer_first": inspect_header_footer_xml(sec.first_page_footer, "footer_first"),
            "footer_even": inspect_header_footer_xml(sec.even_page_footer, "footer_even"),
        }
        hf_data.append(sec_hf)
    return hf_data

def analyze_tables_advanced(doc):
    """Area 3: Advanced table styling (column widths, tblHeader, cantSplit, borders, margins)"""
    tables_info = []
    for idx, table in enumerate(doc.tables):
        tbl_align = get_xml_attr(table._element, './w:tblPr/w:tblAlign', 'w:val')
        
        grid_cols = table._element.xpath('./w:tblGrid/w:gridCol')
        col_widths_pt = [dxa_to_pt(col.get(qn('w:w'))) for col in grid_cols] if grid_cols else []

        tbl_borders = {}
        border_nodes = table._element.xpath('./w:tblPr/w:tblBorders/*')
        for b in border_nodes:
            b_tag = b.tag.split('}')[-1]
            tbl_borders[b_tag] = {
                "val": b.get(qn('w:val')),
                "sz": b.get(qn('w:sz')),
                "space": b.get(qn('w:space')),
                "color": b.get(qn('w:color'))
            }

        tbl_cell_mar = {}
        mar_nodes = table._element.xpath('./w:tblPr/w:tblCellMar/*')
        for m in mar_nodes:
            m_tag = m.tag.split('}')[-1]
            tbl_cell_mar[m_tag] = dxa_to_pt(m.get(qn('w:w')))

        rows_info = []
        is_callout = (len(table.rows) == 1 and len(table.columns) == 1)

        for r_idx, row in enumerate(table.rows):
            has_tbl_header = len(row._tr.xpath('./w:trPr/w:tblHeader')) > 0
            has_cant_split = len(row._tr.xpath('./w:trPr/w:cantSplit')) > 0

            cells_info = []
            for c_idx, cell in enumerate(row.cells):
                shd_elem = cell._tc.xpath('./w:tcPr/w:shd')
                shd_color = shd_elem[0].get(qn('w:fill')) if shd_elem else None
                
                v_align = get_xml_attr(cell._tc, './w:tcPr/w:vAlign', 'w:val')
                grid_span = get_xml_attr(cell._tc, './w:tcPr/w:gridSpan', 'w:val')
                v_merge = get_xml_attr(cell._tc, './w:tcPr/w:vMerge', 'w:val')

                tc_borders = {}
                tc_b_nodes = cell._tc.xpath('./w:tcPr/w:tcBorders/*')
                for cb in tc_b_nodes:
                    cb_tag = cb.tag.split('}')[-1]
                    tc_borders[cb_tag] = {
                        "val": cb.get(qn('w:val')),
                        "sz": cb.get(qn('w:sz')),
                        "color": cb.get(qn('w:color'))
                    }

                cells_info.append({
                    "cell_index": c_idx,
                    "text": cell.text.strip().replace('\n', ' '),
                    "shd_fill": shd_color,
                    "valign": v_align,
                    "grid_span": int(grid_span) if grid_span else 1,
                    "v_merge": v_merge,
                    "cell_borders": tc_borders if tc_borders else None
                })

            rows_info.append({
                "row_index": r_idx,
                "tblHeader": has_tbl_header,
                "cantSplit": has_cant_split,
                "cells": cells_info
            })

        tables_info.append({
            "table_index": idx + 1,
            "is_callout_box": is_callout,
            "rows_count": len(table.rows),
            "cols_count": len(table.columns),
            "table_alignment": tbl_align,
            "column_widths_pt": col_widths_pt,
            "table_borders": tbl_borders,
            "table_cell_margins_pt": tbl_cell_mar,
            "rows": rows_info
        })
    return tables_info

def analyze_typography_nuances(doc):
    """Area 4: Typography nuances (line_spacing, space_before, space_after, keep_with_next, keep_lines)"""
    paragraphs_details = []
    fonts_detected = set()
    colors_detected = set()

    for idx, p in enumerate(doc.paragraphs):
        p_element = p._element
        
        spacing_elem = p_element.xpath('./w:pPr/w:spacing')
        sp_before = dxa_to_pt(spacing_elem[0].get(qn('w:before'))) if spacing_elem else None
        sp_after = dxa_to_pt(spacing_elem[0].get(qn('w:after'))) if spacing_elem else None
        sp_line = spacing_elem[0].get(qn('w:line')) if spacing_elem else None
        sp_line_rule = spacing_elem[0].get(qn('w:lineRule')) if spacing_elem else None

        keep_next = len(p_element.xpath('./w:pPr/w:keepNext')) > 0
        keep_lines = len(p_element.xpath('./w:pPr/w:keepLines')) > 0
        widow_control = len(p_element.xpath('./w:pPr/w:widowControl')) > 0
        page_break_before = len(p_element.xpath('./w:pPr/w:pageBreakBefore')) > 0

        ind_elem = p_element.xpath('./w:pPr/w:ind')
        ind_left = dxa_to_pt(ind_elem[0].get(qn('w:left'))) if ind_elem else None
        ind_right = dxa_to_pt(ind_elem[0].get(qn('w:right'))) if ind_elem else None
        ind_first_line = dxa_to_pt(ind_elem[0].get(qn('w:firstLine'))) if ind_elem else None
        ind_hanging = dxa_to_pt(ind_elem[0].get(qn('w:hanging'))) if ind_elem else None

        jc_elem = p_element.xpath('./w:pPr/w:jc')
        alignment = jc_elem[0].get(qn('w:val')) if jc_elem else None

        runs_info = []
        for r in p.runs:
            r_elem = r._element
            font_name = r.font.name
            if not font_name:
                rFonts = r_elem.xpath('./w:rPr/w:rFonts')
                if rFonts:
                    font_name = rFonts[0].get(qn('w:ascii')) or rFonts[0].get(qn('w:hAnsi'))
            if font_name:
                fonts_detected.add(font_name)

            color_hex = None
            if r.font.color and r.font.color.rgb:
                color_hex = str(r.font.color.rgb)
            else:
                w_color = r_elem.xpath('./w:rPr/w:color')
                if w_color:
                    color_hex = w_color[0].get(qn('w:val'))
            if color_hex:
                colors_detected.add(color_hex)

            sz_elem = r_elem.xpath('./w:rPr/w:sz')
            font_size_pt = float(sz_elem[0].get(qn('w:val'))) / 2.0 if sz_elem else (r.font.size.pt if r.font.size else None)

            runs_info.append({
                "text": r.text,
                "font_name": font_name,
                "font_size_pt": font_size_pt,
                "bold": r.bold,
                "italic": r.italic,
                "color": color_hex
            })

        if (p.style and p.style.name and p.style.name.startswith("Heading")) or p.text.strip():
            paragraphs_details.append({
                "paragraph_index": idx,
                "style": p.style.name if p.style else None,
                "alignment": alignment,
                "spacing": {
                    "space_before_pt": sp_before,
                    "space_after_pt": sp_after,
                    "line_spacing": sp_line,
                    "line_rule": sp_line_rule
                },
                "pagination_controls": {
                    "keep_with_next": keep_next,
                    "keep_lines": keep_lines,
                    "widow_control": widow_control,
                    "page_break_before": page_break_before
                },
                "indents_pt": {
                    "left": ind_left,
                    "right": ind_right,
                    "first_line": ind_first_line,
                    "hanging": ind_hanging
                },
                "runs": runs_info
            })

    return {
        "fonts_detected": sorted(list(fonts_detected)),
        "colors_detected": sorted(list(colors_detected)),
        "sample_paragraphs": paragraphs_details[:30]
    }

def analyze_lists_and_numbering(doc):
    """Area 5: Numbered lists and custom bullet formats (w:numPr, w:abstractNum)"""
    numbering_info = {
        "numbered_paragraphs_count": 0,
        "list_items": [],
        "abstract_numbering_definitions": []
    }

    for idx, p in enumerate(doc.paragraphs):
        num_pr = p._element.xpath('./w:pPr/w:numPr')
        if num_pr:
            numbering_info["numbered_paragraphs_count"] += 1
            num_id_elem = num_pr[0].xpath('./w:numId')
            ilvl_elem = num_pr[0].xpath('./w:ilvl')
            
            num_id = num_id_elem[0].get(qn('w:val')) if num_id_elem else None
            ilvl = ilvl_elem[0].get(qn('w:ilvl')) if ilvl_elem else "0"

            numbering_info["list_items"].append({
                "paragraph_index": idx,
                "style": p.style.name if p.style else None,
                "num_id": num_id,
                "ilvl": int(ilvl),
                "text": p.text.strip()
            })

    try:
        if hasattr(doc.part, 'numbering_part') and doc.part.numbering_part is not None:
            num_part = doc.part.numbering_part.element
            abstract_nums = num_part.xpath('./w:abstractNum')
            for abs_num in abstract_nums:
                abs_id = abs_num.get(qn('w:abstractNumId'))
                levels = []
                for lvl in abs_num.xpath('./w:lvl'):
                    ilvl_val = lvl.get(qn('w:ilvl'))
                    num_fmt = get_xml_attr(lvl, './w:numFmt', 'w:val')
                    lvl_text = get_xml_attr(lvl, './w:lvlText', 'w:val')
                    start_val = get_xml_attr(lvl, './w:start', 'w:val')

                    r_font = get_xml_attr(lvl, './w:rPr/w:rFonts', 'w:ascii')
                    left_ind = dxa_to_pt(get_xml_attr(lvl, './w:pPr/w:ind', 'w:left'))
                    hanging_ind = dxa_to_pt(get_xml_attr(lvl, './w:pPr/w:ind', 'w:hanging'))

                    levels.append({
                        "ilvl": int(ilvl_val) if ilvl_val else 0,
                        "numFmt": num_fmt,
                        "lvlText": lvl_text,
                        "start": start_val,
                        "bullet_font": r_font,
                        "indents_pt": {"left": left_ind, "hanging": hanging_ind}
                    })

                numbering_info["abstract_numbering_definitions"].append({
                    "abstractNumId": abs_id,
                    "levels": levels
                })
    except Exception as e:
        numbering_info["numbering_xml_error"] = str(e)

    return numbering_info

def analyze_xml_sdt_and_fields(doc):
    sdts = doc._element.body.xpath('.//*[local-name()="sdt"]')
    sdt_details = []
    for sdt in sdts:
        alias = get_xml_attr(sdt, './w:sdtPr/w:alias', 'w:val')
        tag = get_xml_attr(sdt, './w:sdtPr/w:tag', 'w:val')
        is_toc = len(sdt.xpath('./w:sdtPr/w:docPartObj/w:docPartGallery[@w:val="Table of Contents"]')) > 0
        sdt_details.append({
            "alias": alias,
            "tag": tag,
            "is_toc_sdt": is_toc
        })

    instr_texts = doc._element.body.xpath('.//*[local-name()="instrText"]')
    field_codes = [ (t.text or "").strip() for t in instr_texts if (t.text or "").strip() ]

    return {
        "sdt_blocks_count": len(sdts),
        "sdt_details": sdt_details,
        "field_codes_count": len(field_codes),
        "field_codes": field_codes,
        "has_toc_field": any("TOC" in c for c in field_codes)
    }

def analyze_docx_file(filepath):
    if not os.path.exists(filepath):
        print(json.dumps({"error": f"File {filepath} not found."}, ensure_ascii=False))
        return

    doc = docx.Document(filepath)

    sections_info = analyze_sections(doc)
    headers_footers_info = analyze_headers_footers(doc)
    tables_info = analyze_tables_advanced(doc)
    typography_info = analyze_typography_nuances(doc)
    lists_info = analyze_lists_and_numbering(doc)
    xml_info = analyze_xml_sdt_and_fields(doc)

    analysis = {
        "filepath": filepath,
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "total_sections": len(doc.sections),
        "section_properties": sections_info,
        "headers_and_footers": headers_footers_info,
        "table_grid_specs": tables_info,
        "typography_and_paragraph_rules": typography_info,
        "list_and_numbering_rules": lists_info,
        "xml_sdt_and_fields": xml_info
    }

    print(json.dumps(analysis, indent=2, ensure_ascii=False))

if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else "MAU REPORT.docx"
    analyze_docx_file(target)
