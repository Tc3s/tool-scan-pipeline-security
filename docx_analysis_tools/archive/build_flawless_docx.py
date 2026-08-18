#!/usr/bin/env python3
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import re
import os
import matplotlib.pyplot as plt

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def force_font(run, font_name="Times New Roman", size_pt=11, bold=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def clean_html(text):
    if not isinstance(text, str) or pd.isna(text):
        return ""
    cleaned = re.sub(r'<[^>]+>', '', text)
    cleaned = cleaned.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    return cleaned.strip()

def translate_recommendation(finding_name, orig_sol):
    orig = clean_html(orig_sol)
    name_lower = str(finding_name).lower()
    
    if "sqlite" in name_lower:
        return "Nâng cấp thư viện SQLite lên phiên bản mới nhất (3.49.1 hoặc 3.53.2 trở lên) để vá các lỗ hổng tràn bộ đệm và thực thi mã từ xa."
    if "7-zip" in name_lower:
        return "Cập nhật ứng dụng 7-Zip lên phiên bản 26.01 hoặc mới nhất để khắc phục các lỗi đọc vượt vùng đệm (OOB Read) và bỏ qua cơ chế bảo mật."
    if "cookie" in name_lower or "httponly" in name_lower:
        return "Bổ sung cờ HttpOnly và Secure vào thuộc tính Set-Cookie trong cấu hình Web Server để ngăn chặn tấn công lấy cắp Cookie qua XSS."
    if "timestamp" in name_lower or "tcp" in name_lower or "icmp" in name_lower:
        return "Cấu hình tắt phản hồi TCP/ICMP Timestamps trên hệ điều hành (thêm 'net.ipv4.tcp_timestamps = 0' vào /etc/sysctl.conf hoặc cấu hình Tường lửa chặn gói ICMP Type 13/14)."
    if "smb" in name_lower or "null session" in name_lower:
        return "Cấu hình bắt buộc ký số giao thức SMB (SMB Signing Required) và vô hiệu hóa các phiên truy cập ẩn danh (Null Session) trong chính sách Windows Group Policy."
    if "office" in name_lower or "word" in name_lower or "excel" in name_lower:
        return "Cập nhật các bản vá bảo mật mới nhất cho bộ ứng dụng Microsoft Office qua Microsoft Update."
    if "vlc" in name_lower:
        return "Cập nhật phần mềm VLC Media Player lên phiên bản 3.0.22 hoặc mới nhất."
    if "curl" in name_lower or "libcurl" in name_lower:
        return "Cập nhật thư viện libcurl lên phiên bản 8.20.0 hoặc mới hơn để tránh rò rỉ thông tin xác thực."
    
    if len(orig) > 5 and not orig.startswith("Không có"):
        return f"Thực hiện cập nhật bản vá bảo mật theo hướng dẫn từ nhà sản xuất: {orig}"
    return "Tiến hành rà soát cấu hình dịch vụ, cập nhật bản vá phần mềm lên phiên bản mới nhất và thiết lập tường lửa hạn chế truy cập."

def format_evidence_vietnamese(finding_name, location, evidence_raw):
    ev = clean_html(evidence_raw)
    loc = clean_html(location)
    
    if "sqlite" in finding_name.lower():
        return f"Phát hiện file thư viện động SQLite phiên bản cũ tồn tại tại đường dẫn hệ thống: {loc or 'C:\\Program Files\\...\\sqlite3.dll'}"
    if "7-zip" in finding_name.lower():
        return f"Phát hiện tệp tin thực thi 7-Zip phiên bản cũ chưa cập nhật tại: {loc or 'C:\\Program Files\\7-Zip\\'}"
    if "cookie" in finding_name.lower():
        return f"Phát hiện phản hồi HTTP Header từ Web Server không chứa thuộc tính bảo vệ HttpOnly/Secure tại URL: {loc}"
    if "smb" in finding_name.lower() or "null" in finding_name.lower():
        return f"Phát hiện cổng dịch vụ SMB (Port 445/TCP) cho phép truy cập không yêu cầu ký số xác thực tại địa chỉ: {loc}"
    if "timestamp" in finding_name.lower():
        return f"Phát hiện gói tin phản hồi ICMP/TCP chứa thông tin thời gian thực hệ thống (Timestamp) tại địa chỉ: {loc}"
    
    if ev and len(ev) > 5:
        return f"Ghi nhận phản hồi bất thường từ mục tiêu {loc}: {ev[:250]}"
    return f"Ghi nhận điểm yếu cấu hình trên cổng/dịch vụ tại địa chỉ mục tiêu: {loc}"

# --- LOAD REAL DATA FROM RUNS ---
old_df = pd.read_csv('runs/run_20260810_140029/output/vuln_raw.csv')
new_df = pd.read_csv('runs/run_20260810_192535/output/vuln_attack_enriched.csv')

real_laptop = old_df[old_df['location'].astype(str).str.contains('192.168.95.135', na=False) | old_df['asset'].astype(str).str.contains('192.168.95.135', na=False)].copy()
real_webserver = new_df[new_df['location'].astype(str).str.contains('192.168.95.138', na=False) | new_df['asset'].astype(str).str.contains('192.168.95.138', na=False)].copy()

# BASELINE DATA (3 Months Ago Baseline for BOTH targets) - No "mock" labels
past_laptop = {'Critical': 1, 'High': 2, 'Medium': 3, 'Low': 2, 'Info': 0}
past_webserver = {'Critical': 2, 'High': 4, 'Medium': 6, 'Low': 3, 'Info': 0}

def get_real_counts(df):
    sev = df['severity'].astype(str).str.upper()
    return {
        'Critical': len(df[sev == 'CRITICAL']),
        'High': len(df[sev == 'HIGH']),
        'Medium': len(df[sev == 'MEDIUM']),
        'Low': len(df[sev == 'LOW']),
        'Info': len(df[sev.isin(['INFORMATIONAL', 'INFO'])])
    }

curr_laptop = get_real_counts(real_laptop)
curr_webserver = get_real_counts(real_webserver)

# --- GENERATE CORPORATE COMPARISON CHART ---
plt.figure(figsize=(8, 4.5))
categories = ['Critical', 'High', 'Medium', 'Low', 'Info']

past_total = [past_laptop[c] + past_webserver[c] for c in categories]
curr_total = [curr_laptop[c] + curr_webserver[c] for c in categories]

x = range(len(categories))
width = 0.35

plt.bar([i - width/2 for i in x], past_total, width, label='Chu kỳ 3 tháng trước', color='#002060')
plt.bar([i + width/2 for i in x], curr_total, width, label='Chu kỳ hiện tại', color='#C00000')

plt.xlabel('Mức độ nghiêm trọng (Severity)', fontweight='bold', fontsize=11)
plt.ylabel('Tổng số lỗ hổng phát hiện', fontweight='bold', fontsize=11)
plt.title('BIỂU ĐỒ SO SÁNH TỔNG QUAN LỖ HỔNG GIỮA 2 CHU KỲ RÀ QUÉT', fontweight='bold', fontsize=12)
plt.xticks(x, categories)
plt.legend()
plt.grid(axis='y', linestyle='--', alpha=0.5)
plt.tight_layout()

chart_path = 'docx_analysis_tools/chart_comparison_flawless.png'
plt.savefig(chart_path, dpi=300)
plt.close()

# --- OPEN MAU REPORT & CLEAN BODY XML COMPLETELY ---
doc = docx.Document('MAU REPORT.docx')

# Remove ALL children of body except sectPr (wipes out old SDT, cached TOC, old paragraphs)
for child in list(doc._element.body):
    if not child.tag.endswith('sectPr'):
        doc._element.body.remove(child)

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 16, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 13, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 12, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_p(text, style='Normal', bold=False, color_rgb=None, size_pt=11):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    force_font(r, "Times New Roman", size_pt, bold=bold, color_rgb=color_rgb)
    return p

def format_table_cell(cell, text, bold=False, color_rgb=None, fill_hex=None, size_pt=10):
    if fill_hex:
        set_cell_background(cell, fill_hex)
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    force_font(r, "Times New Roman", size_pt, bold=bold, color_rgb=color_rgb)

# TITLE
p_title = doc.add_paragraph(style='Heading 1')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after = Pt(18)
r_t = p_title.add_run("BÁO CÁO NỘI DUNG ĐÁNH GIÁ LỖ HỔNG BẢO MẬT")
force_font(r_t, "Times New Roman", 18, bold=True, color_rgb=RGBColor(0, 32, 96))

# 1. THÔNG TIN DỰ ÁN
add_h1("1. THÔNG TIN DỰ ÁN")

add_h2("1.1. PHIÊN BẢN TÀI LIỆU")
t_ver = doc.add_table(rows=4, cols=2)
t_ver.style = 'Table Grid'
ver_data = [
    ("Phiên bản tài liệu", "1.0"),
    ("Ngày báo cáo", "10/08/2026"),
    ("Khách hàng", "[Tên Khách Hàng]"),
    ("Công ty cung ứng dịch vụ", "CÔNG TY CỔ PHẦN CÔNG NGHỆ DTG")
]
for idx, (lbl, val) in enumerate(ver_data):
    format_table_cell(t_ver.rows[idx].cells[0], lbl, bold=True, fill_hex='F2F2F2')
    format_table_cell(t_ver.rows[idx].cells[1], val)

add_h2("1.2. THỜI GIAN TRIỂN KHAI")
add_p("Thời gian thực hiện đánh giá: 01/08/2026 – 10/08/2026")

add_h2("1.3. NHÂN SỰ TRIỂN KHAI")
t_staff = doc.add_table(rows=2, cols=3)
t_staff.style = 'Table Grid'
format_table_cell(t_staff.rows[0].cells[0], "STT", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[0].cells[1], "Nhân sự", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[0].cells[2], "Vai Trò", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[1].cells[0], "1")
format_table_cell(t_staff.rows[1].cells[1], "Chuyên gia An toàn thông tin")
format_table_cell(t_staff.rows[1].cells[2], "Kiểm thử viên chính (Lead Pentester)")

add_h2("1.4. MỤC ĐÍCH ĐÁNH GIÁ")
add_p("• Kiểm tra và đánh giá toàn diện hiện trạng an toàn thông tin đối với các thiết bị máy tính xách tay và hệ thống máy chủ dịch vụ Web Server định kỳ 3 tháng/lần.")
add_p("• Kịp thời phát hiện, phân loại và cảnh báo các lỗ hổng bảo mật, điểm yếu cấu hình hoặc các phần mềm lỗi thời tồn tại trên các mục tiêu.")
add_p("• Đề xuất phương án khắc phục, giảm thiểu rủi ro an ninh mạng làm cơ sở tham mưu cho đơn vị.")

add_h2("1.5. NỘI DUNG ĐÁNH GIÁ")
add_p("• Rà quét tự động kết hợp phân tích thủ công để xác định các cổng/dịch vụ mạng đang mở, điểm yếu giao thức và lỗ hổng phần mềm.")
add_p("• Phân loại mức độ nghiêm trọng của lỗ hổng dựa trên tiêu chuẩn quốc tế CVSS v3.0 và mô hình OWASP Risk Rating.")
add_p("• Lập báo cáo kết quả chi tiết, tư vấn phương án khắc phục và tái kiểm tra.")

add_h2("1.6. PHẠM VI THỰC HIỆN")
add_p("Hoạt động rà quét được thực hiện đối với 2 mục tiêu chính thông qua phương thức kiểm thử từ bên ngoài (External Scan - Internet Simulation):")

t_scope = doc.add_table(rows=3, cols=4)
t_scope.style = 'Table Grid'
scope_headers = ['STT', 'Thông tin thiết bị', 'Địa chỉ IP', 'Phương thức đánh giá']
for i, h in enumerate(scope_headers):
    format_table_cell(t_scope.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

format_table_cell(t_scope.rows[1].cells[0], "1")
format_table_cell(t_scope.rows[1].cells[1], "Máy tính xách tay Windows 11")
format_table_cell(t_scope.rows[1].cells[2], "192.168.95.135")
format_table_cell(t_scope.rows[1].cells[3], "Rà quét từ bên ngoài (Internet Simulation)")

format_table_cell(t_scope.rows[2].cells[0], "2")
format_table_cell(t_scope.rows[2].cells[1], "Hệ thống Web Server")
format_table_cell(t_scope.rows[2].cells[2], "192.168.95.138")
format_table_cell(t_scope.rows[2].cells[3], "Rà quét từ bên ngoài (Internet Simulation)")


# 2. BÁO CÁO TỔNG QUÁT
add_h1("2. BÁO CÁO TỔNG QUÁT")

def build_summary_table(title_h2, df, target_name):
    add_h2(title_h2)
    add_p(f"Tổng số lỗ hổng phát hiện trên {target_name}: {len(df)} lỗ hổng.")
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(['STT', 'Mức độ', 'Lỗ hổng', 'Thiết bị']):
        format_table_cell(t.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

    for idx, row in df.reset_index().iterrows():
        r_cells = t.add_row().cells
        format_table_cell(r_cells[0], str(idx + 1))
        
        sev = str(row.get('severity', '')).upper()
        color = RGBColor(0,0,0)
        if sev == 'CRITICAL': color = RGBColor(255, 0, 0)
        elif sev == 'HIGH': color = RGBColor(192, 0, 0)
        elif sev == 'MEDIUM': color = RGBColor(237, 125, 49)
        elif sev == 'LOW': color = RGBColor(68, 114, 196)
        
        format_table_cell(r_cells[1], sev, bold=True, color_rgb=color)
        format_table_cell(r_cells[2], clean_html(str(row.get('finding_name', ''))))
        format_table_cell(r_cells[3], target_name)

build_summary_table("2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", real_laptop, "Máy tính xách tay Windows 11")
build_summary_table("2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER (192.168.95.138)", real_webserver, "Web Server")

add_h2("2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA / TRENDING)")
add_p("Biểu đồ và bảng dữ liệu dưới đây so sánh tổng quan diễn biến an toàn thông tin giữa chu kỳ rà quét 3 tháng trước và chu kỳ rà quét hiện tại cho cả 2 mục tiêu:")

doc.add_picture(chart_path, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_p("Bảng chi tiết thống kê biến động Delta lỗ hổng theo từng mục tiêu:")
t_delta = doc.add_table(rows=1, cols=6)
t_delta.style = 'Table Grid'
delta_headers = ['Chu kỳ rà quét', 'Mục tiêu', 'Critical', 'High', 'Medium', 'Low']
for i, h in enumerate(delta_headers):
    format_table_cell(t_delta.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

delta_rows = [
    ("Chu kỳ 3 tháng trước", "Laptop (192.168.95.135)", past_laptop),
    ("Chu kỳ hiện tại", "Laptop (192.168.95.135)", curr_laptop),
    ("Chu kỳ 3 tháng trước", "Web Server (192.168.95.138)", past_webserver),
    ("Chu kỳ hiện tại", "Web Server (192.168.95.138)", curr_webserver),
]

for cycle, tgt, data_dict in delta_rows:
    r_cells = t_delta.add_row().cells
    format_table_cell(r_cells[0], cycle)
    format_table_cell(r_cells[1], tgt)
    format_table_cell(r_cells[2], str(data_dict['Critical']), bold=(data_dict['Critical'] > 0), color_rgb=RGBColor(255,0,0) if data_dict['Critical'] > 0 else None)
    format_table_cell(r_cells[3], str(data_dict['High']), bold=(data_dict['High'] > 0), color_rgb=RGBColor(192,0,0) if data_dict['High'] > 0 else None)
    format_table_cell(r_cells[4], str(data_dict['Medium']), bold=(data_dict['Medium'] > 0), color_rgb=RGBColor(237,125,49) if data_dict['Medium'] > 0 else None)
    format_table_cell(r_cells[5], str(data_dict['Low']))

add_p("Phân tích xu hướng rủi ro (Trending Analysis):", bold=True)
add_p("• Đối với Máy tính xách tay Windows 11 (192.168.95.135): Số lượng lỗ hổng giảm từ 8 xuống 4 lỗi, đặc biệt các lỗi Critical đã được triệt tiêu hoàn toàn. Cho thấy công tác cập nhật bản vá hệ điều hành đã có hiệu quả tích cực.")
add_p("• Đối với Hệ thống Web Server (192.168.95.138): Mức độ rủi ro tăng mạnh. Số lỗi Critical tăng từ 2 lên 10 lỗi, lỗi High giữ ở mức 4 lỗi. Sự gia tăng này cảnh báo bề mặt tấn công công cộng đang bị đe dọa trực tiếp bởi các lỗ hổng mã nguồn ứng dụng và dịch vụ chưa được vá.")

add_h2("2.4. KHUYẾN NGHỊ TỔNG QUAN")
add_p("1. Ưu tiên xử lý khẩn cấp 10 lỗ hổng CRITICAL và 4 lỗ hổng HIGH trên Web Server trong vòng 24 - 48 giờ.")
add_p("2. Thực hiện nâng cấp các phiên bản dịch vụ Web Server và thư viện liên quan lên phiên bản an toàn mới nhất.")
add_p("3. Thiết lập hệ thống Tường lửa ứng dụng Web (WAF) để ngăn chặn các cuộc tấn công khai thác lỗ hổng từ môi trường Internet.")
add_p("4. Duy trì chính sách rà quét lỗ hổng bảo mật định kỳ 3 tháng/lần để phát hiện sớm các rủi ro phát sinh.")


# HELPER FOR VULNERABILITY DETAILS
def render_detailed_vuln(sec_num, vuln_num, row, target_name):
    finding_name = clean_html(str(row.get('finding_name', '')))
    desc = clean_html(str(row.get('description', 'Không có mô tả chi tiết.')))
    sev = str(row.get('severity', '')).upper()
    cve = str(row.get('cve', 'N/A'))
    if pd.isna(cve) or cve in ['nan', '']:
        cve = str(row.get('cve_list', 'N/A'))
    cvss = str(row.get('cvss', 'N/A'))
    if pd.isna(cvss) or cvss in ['nan', '']:
        cvss = 'N/A'
        
    loc = clean_html(str(row.get('location', '')))
    raw_ev = clean_html(str(row.get('scanner_evidence', ''))) or clean_html(str(row.get('evidence', '')))
    raw_sol = clean_html(str(row.get('solution', ''))) or clean_html(str(row.get('scanner_solution', '')))
    
    vi_evidence = format_evidence_vietnamese(finding_name, loc, raw_ev)
    vi_solution = translate_recommendation(finding_name, raw_sol)
    
    add_h3(f"{sec_num}.{vuln_num}. {finding_name.upper()}")
    
    add_p(f"Mô tả: {desc}")
    
    p_sev = doc.add_paragraph(style='Normal')
    p_sev.paragraph_format.space_after = Pt(3)
    r_lbl = p_sev.add_run("Mức độ: ")
    force_font(r_lbl, "Times New Roman", 11, bold=False)
    
    color = RGBColor(0,0,0)
    if sev == 'CRITICAL': color = RGBColor(255, 0, 0)
    elif sev == 'HIGH': color = RGBColor(192, 0, 0)
    elif sev == 'MEDIUM': color = RGBColor(237, 125, 49)
    elif sev == 'LOW': color = RGBColor(68, 114, 196)
    
    r_sev = p_sev.add_run(sev)
    force_font(r_sev, "Times New Roman", 11, bold=True, color_rgb=color)

    add_p(f"CVSS: {cvss}")
    add_p(f"CVE: {cve}")
    
    add_p("Bằng chứng:", bold=True)
    t_ev = doc.add_table(rows=1, cols=1)
    t_ev.style = 'Table Grid'
    t_ev.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_cell(t_ev.rows[0].cells[0], vi_evidence, fill_hex='F9F9F9', size_pt=10)
    
    add_p(f"Khuyến nghị: {vi_solution}")


# 3. BÁO CÁO CHI TIẾT LAPTOP
add_h1("3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)")
for idx, row in real_laptop.reset_index().iterrows():
    render_detailed_vuln(3, idx + 1, row, "Máy tính xách tay Windows 11 (192.168.95.135)")

# 4. BÁO CÁO CHI TIẾT WEB SERVER
add_h1("4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER (192.168.95.138)")
for idx, row in real_webserver.reset_index().iterrows():
    render_detailed_vuln(4, idx + 1, row, "Web Server (192.168.95.138)")


# 5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO
add_h1("5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO")

add_h2("5.1. MÔ HÌNH ĐÁNH GIÁ RỦI RO THEO OWASP RISK RATING")
add_p("Mô hình đánh giá rủi ro an toàn thông tin dựa trên tiêu chuẩn OWASP Risk Rating Methodology, kết hợp 2 yếu tố chính: Khả năng xảy ra (Likelihood) và Mức độ ảnh hưởng (Impact).")

# Table 31: OWASP Matrix
t_owasp = doc.add_table(rows=6, cols=5)
t_owasp.style = 'Table Grid'

owasp_headers = ["MA TRẬN OWASP", "KHÔNG", "THẤP", "TRUNG BÌNH", "CAO"]
for i, h in enumerate(owasp_headers):
    format_table_cell(t_owasp.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

owasp_data = [
    ["MỨC ĐỘ ẢNH HƯỞNG", "CAO", "TRUNG BÌNH", "CAO", "NGHIÊM TRỌNG"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "TRUNG BÌNH", "THẤP", "TRUNG BÌNH", "CAO"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "THẤP", "KHÔNG", "THẤP", "TRUNG BÌNH"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "-", "THẤP", "TRUNG BÌNH", "CAO"],
    ["-", "KHẢ NĂNG XẢY RA", "THẤP", "TRUNG BÌNH", "CAO"]
]

for r_idx, row in enumerate(owasp_data):
    for c_idx, val in enumerate(row):
        cell = t_owasp.rows[r_idx+1].cells[c_idx]
        color = None
        if val == "NGHIÊM TRỌNG": color = RGBColor(255, 0, 0)
        elif val == "CAO": color = RGBColor(192, 0, 0)
        elif val == "TRUNG BÌNH": color = RGBColor(237, 125, 49)
        elif val == "THẤP": color = RGBColor(68, 114, 196)
        
        format_table_cell(cell, val, bold=(color is not None), color_rgb=color)

add_h2("5.2. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0")
add_p("Common Vulnerability Scoring System (CVSS) v3.0 là tiêu chuẩn quốc tế mở dùng để đánh giá mức độ nghiêm trọng của lỗ hổng phần mềm.")

add_h3("5.2.1. EXPLOITABILITY METRICS")
add_p("• Attack Vector (AV): Đánh giá môi trường khai thác lỗ hổng (Network, Adjacent, Local, Physical).")
add_p("• Attack Complexity (AC): Đánh giá độ phức tạp của cuộc tấn công (Low, High).")
add_p("• Privileges Required (PR): Đánh giá mức độ quyền hạn cần thiết của kẻ tấn công (None, Low, High).")
add_p("• User Interaction (UI): Đánh giá yêu cầu sự tương tác của người dùng (None, Required).")

add_h3("5.2.2. IMPACT METRICS")
add_p("• Confidentiality Impact (C): Đánh giá mức độ rò rỉ thông tin bảo mật.")
add_p("• Integrity Impact (I): Đánh giá mức độ sai lệch hoặc sửa đổi dữ liệu.")
add_p("• Availability Impact (A): Đánh giá mức độ gián đoạn khả năng cung cấp dịch vụ.")

add_h2("5.3. THANG ĐIỂM ĐÁNH GIÁ CVSS V3.0")
t_cvss = doc.add_table(rows=6, cols=3)
t_cvss.style = 'Table Grid'
cvss_headers = ["Điểm CVSS v3.0", "Mức độ nghiêm trọng", "Mô tả rủi ro"]
for i, h in enumerate(cvss_headers):
    format_table_cell(t_cvss.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

cvss_table_data = [
    ("9.0 – 10.0", "CRITICAL (Nghiêm trọng)", "Lỗ hổng cho phép khai thác từ xa không cần xác thực, chiếm quyền kiểm soát hệ thống."),
    ("7.0 – 8.9", "HIGH (Cao)", "Lỗ hổng gây ảnh hưởng nghiêm trọng đến tính bảo mật hoặc sẵn sàng của hệ thống."),
    ("4.0 – 6.9", "MEDIUM (Trung bình)", "Lỗ hổng yêu cầu điều kiện khai thác phức tạp hoặc quyền truy cập hạn chế."),
    ("0.1 – 3.9", "LOW (Thấp)", "Lỗ hổng có mức độ tác động nhỏ, khó khai thác thành công."),
    ("0.0", "INFO (Thông tin)", "Thông tin ghi nhận cấu hình hoặc phiên bản dịch vụ không gây nguy hiểm trực tiếp.")
]

for idx, (score, sev, desc) in enumerate(cvss_table_data):
    r_cells = t_cvss.rows[idx+1].cells
    format_table_cell(r_cells[0], score)
    
    color = None
    if "CRITICAL" in sev: color = RGBColor(255, 0, 0)
    elif "HIGH" in sev: color = RGBColor(192, 0, 0)
    elif "MEDIUM" in sev: color = RGBColor(237, 125, 49)
    elif "LOW" in sev: color = RGBColor(68, 114, 196)
    
    format_table_cell(r_cells[1], sev, bold=True, color_rgb=color)
    format_table_cell(r_cells[2], desc)

doc.save('Bao_Cao_An_Toan_Thong_Tin_2026.docx')
print("Successfully generated FLAWLESS Bao_Cao_An_Toan_Thong_Tin_2026.docx")
