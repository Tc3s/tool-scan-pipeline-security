#!/usr/bin/env python3
import docx
from docx.shared import Pt
import json
import os

def analyze_docx(filepath):
    if not os.path.exists(filepath):
        print(f"Error: {filepath} not found.")
        return

    doc = docx.Document(filepath)
    
    analysis = {
        "styles": set(),
        "fonts": set(),
        "colors": set(),
        "has_images": len(doc.inline_shapes) > 0,
        "tables_count": len(doc.tables),
        "text_sample": [],
        "headings": [],
        "sections": len(doc.sections)
    }
    
    for para in doc.paragraphs:
        if para.style and para.style.name:
            analysis["styles"].add(para.style.name)
            
        if para.text.strip():
            if 'Heading' in para.style.name:
                analysis["headings"].append(para.text.strip())
            elif len(analysis["text_sample"]) < 15:
                analysis["text_sample"].append(para.text.strip())
                
        for run in para.runs:
            if run.font.name:
                analysis["fonts"].add(run.font.name)
            if run.font.color and run.font.color.rgb:
                analysis["colors"].add(str(run.font.color.rgb))
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name:
                            analysis["fonts"].add(run.font.name)

    # Convert sets to lists for JSON serialization
    analysis["styles"] = list(analysis["styles"])
    analysis["fonts"] = list(analysis["fonts"])
    analysis["colors"] = list(analysis["colors"])
    
    print(json.dumps(analysis, indent=2, ensure_ascii=False))

if __name__ == "__main__":
    analyze_docx("MAU REPORT.docx")
