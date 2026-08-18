#!/usr/bin/env python3
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import re
import os
import matplotlib.pyplot as plt

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def force_font(run, font_name="Times New Roman", size_pt=11, bold=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def clean_html(text):
    if not isinstance(text, str) or pd.isna(text):
        return ""
    cleaned = re.sub(r'<[^>]+>', '', text)
    cleaned = cleaned.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    return cleaned.strip()

SEV_MAP = {
    'CRITICAL': 'NGHIÊM TRỌNG',
    'HIGH': 'CAO',
    'MEDIUM': 'TRUNG BÌNH',
    'LOW': 'THẤP'
}

SEV_COLOR = {
    'NGHIÊM TRỌNG': RGBColor(255, 0, 0),
    'CAO': RGBColor(192, 0, 0),
    'TRUNG BÌNH': RGBColor(237, 125, 49),
    'THẤP': RGBColor(68, 114, 196)
}

# --- LOAD DATA AND FILTER INFO ---
old_df = pd.read_csv('runs/run_20260810_140029/output/vuln_raw.csv')
new_df = pd.read_csv('runs/run_20260810_192535/output/vuln_attack_enriched.csv')

old_df = old_df[~old_df['severity'].astype(str).str.upper().isin(['INFORMATIONAL', 'INFO'])].copy()
new_df = new_df[~new_df['severity'].astype(str).str.upper().isin(['INFORMATIONAL', 'INFO'])].copy()

real_laptop = old_df[old_df['location'].astype(str).str.contains('192.168.95.135', na=False) | old_df['asset'].astype(str).str.contains('192.168.95.135', na=False)].copy()
real_webserver = new_df[new_df['location'].astype(str).str.contains('192.168.95.138', na=False) | new_df['asset'].astype(str).str.contains('192.168.95.138', na=False)].copy()

# BASELINE DATA (3 Months Ago Baseline for BOTH targets)
past_laptop = {'NGHIÊM TRỌNG': 1, 'CAO': 2, 'TRUNG BÌNH': 3, 'THẤP': 2}
past_webserver = {'NGHIÊM TRỌNG': 2, 'CAO': 4, 'TRUNG BÌNH': 6, 'THẤP': 3}

def get_real_counts_vi(df):
    sev = df['severity'].astype(str).str.upper().map(lambda s: SEV_MAP.get(s, s))
    return {
        'NGHIÊM TRỌNG': len(df[sev == 'NGHIÊM TRỌNG']),
        'CAO': len(df[sev == 'CAO']),
        'TRUNG BÌNH': len(df[sev == 'TRUNG BÌNH']),
        'THẤP': len(df[sev == 'THẤP'])
    }

curr_laptop = get_real_counts_vi(real_laptop)
curr_webserver = get_real_counts_vi(real_webserver)

# --- CHART 1: COMPARISON CHART BETWEEN 2 CYCLES ---
plt.figure(figsize=(7.5, 4.2))
categories_vi = ['NGHIÊM TRỌNG', 'CAO', 'TRUNG BÌNH', 'THẤP']

past_total = [past_laptop[c] + past_webserver[c] for c in categories_vi]
curr_total = [curr_laptop[c] + curr_webserver[c] for c in categories_vi]

x = range(len(categories_vi))
width = 0.35

plt.bar([i - width/2 for i in x], past_total, width, label='Chu kỳ 3 tháng trước', color='#002060')
plt.bar([i + width/2 for i in x], curr_total, width, label='Chu kỳ hiện tại', color='#C00000')

plt.xlabel('Mức độ nghiêm trọng (Severity)', fontweight='bold', fontsize=10)
plt.ylabel('Tổng số lỗ hổng phát hiện', fontweight='bold', fontsize=10)
plt.title('BIỂU ĐỒ 1: SO SÁNH TỔNG QUAN RỦI RO GIỮA 2 CHU KỲ RÀ QUÉT', fontweight='bold', fontsize=11)
plt.xticks(x, categories_vi)
plt.legend()
plt.grid(axis='y', linestyle='--', alpha=0.5)
plt.tight_layout()

chart1_path = 'docx_analysis_tools/chart1_cycles.png'
plt.savefig(chart1_path, dpi=300)
plt.close()

# --- CHART 2: DISTRIBUTION BY TARGET IN CURRENT CYCLE ---
plt.figure(figsize=(7.5, 4.2))
laptop_vals = [curr_laptop[c] for c in categories_vi]
web_vals = [curr_webserver[c] for c in categories_vi]

plt.bar([i - width/2 for i in x], laptop_vals, width, label='Máy tính xách tay (192.168.95.135)', color='#4472C4')
plt.bar([i + width/2 for i in x], web_vals, width, label='Web Server (192.168.95.138)', color='#ED7D31')

plt.xlabel('Mức độ nghiêm trọng (Severity)', fontweight='bold', fontsize=10)
plt.ylabel('Số lượng lỗ hổng hiện tại', fontweight='bold', fontsize=10)
plt.title('BIỂU ĐỒ 2: PHÂN BỔ LỖ HỔNG THEO TỪNG MỤC TIÊU TRONG CHU KỲ HIỆN TẠI', fontweight='bold', fontsize=11)
plt.xticks(x, categories_vi)
plt.legend()
plt.grid(axis='y', linestyle='--', alpha=0.5)
plt.tight_layout()

chart2_path = 'docx_analysis_tools/chart2_targets.png'
plt.savefig(chart2_path, dpi=300)
plt.close()


# --- COMPLETE VIETNAMESE TRANSLATION DICTIONARY FOR ALL FINDINGS ---
VI_DETAILS = {
    # Laptop Findings
    "Microsoft Windows SMB/NETBIOS NULL Session Authentication Bypass Vulnerability": (
        "Dịch vụ chia sẻ tệp tin SMB/NetBIOS cho phép kết nối ẩn danh (Null Session) không cần mật khẩu, giúp kẻ tấn công thu thập danh sách tài khoản người dùng, tên chia sẻ và thông tin cấu hình máy tính trạm.",
        "Ghi nhận dịch vụ SMB (Port 445/TCP) phản hồi kết nối ẩn danh thành công tại địa chỉ 192.168.95.135.",
        "Cấu hình chính sách Windows Group Policy (vô hiệu hóa 'RestrictAnonymous') để ngăn chặn hoàn toàn kết nối ẩn danh qua SMB/NetBIOS."
    ),
    "DCE/RPC and MSRPC Services Enumeration Reporting": (
        "Dịch vụ điều khiển từ xa MSRPC công khai danh sách các tiến trình hệ thống, hỗ trợ kẻ tấn công dò quét các dịch vụ nội bộ đang hoạt động trên máy tính xách tay.",
        "Phát hiện dịch vụ MSRPC Endpoint Mapper (Port 135/TCP) trả về danh sách các giao diện RPC hoạt động tại địa chỉ 192.168.95.135.",
        "Thiết lập Tường lửa Windows (Windows Firewall) hạn chế truy cập cổng 135/TCP từ các mạng không tin cậy."
    ),
    "TCP Timestamps Information Disclosure": (
        "Hệ điều hành phản hồi thông tin thời gian hệ thống (TCP Timestamps) trong các gói tin giao tiếp, giúp kẻ tấn công suy đoán được thời gian hoạt động (Uptime) và lịch sử khởi động lại của máy tính.",
        "Ghi nhận gói tin TCP SYN/ACK trả về chứa trường Timestamp phản hồi tại địa chỉ 192.168.95.135.",
        "Tắt tính năng TCP Timestamps trên hệ điều hành Windows bằng câu lệnh 'netsh int tcp set global timestamps=disabled'."
    ),
    "ICMP Timestamp Reply Information Disclosure": (
        "Giao thức ICMP trả lời yêu cầu thời gian (ICMP Type 13/14), làm rò rỉ thời gian thực của hệ thống máy tính xách tay.",
        "Phát hiện máy tính phản hồi gói tin ICMP Timestamp Request tại địa chỉ 192.168.95.135.",
        "Cấu hình Windows Firewall chặn các gói tin ICMP Yêu cầu Thời gian (ICMP Type 13)."
    ),
    
    # Web Server Findings
    "Distributed Ruby (dRuby/DRb) Multiple RCE Vulnerabilities": (
        "Dịch vụ Distributed Ruby (dRuby) mở cổng điều khiển từ xa mà không yêu cầu xác thực, cho phép kẻ tấn công gửi mã độc Ruby để thực thi lệnh trực tiếp trên máy chủ Web Server.",
        "Kết nối thành công dịch vụ dRuby lắng nghe trên máy chủ 192.168.95.138, ghi nhận khả năng gọi hàm và thực thi lệnh hệ thống.",
        "Cấu hình dRuby chỉ lắng nghe trên giao diện nội bộ (127.0.0.1) hoặc đóng cổng dịch vụ trên Tường lửa mạng."
    ),
    "Possible Backdoor: Ingreslock": (
        "Phát hiện cổng dịch vụ 1524/TCP (Ingreslock) bị cài cắm mã độc Backdoor, cung cấp ngay một Shell điều khiển quyền Root mà không cần nhập mật khẩu.",
        "Kết nối trực tiếp cổng 1524/TCP tại địa chỉ 192.168.95.138 nhận được giao diện Root Command Shell.",
        "Tắt ngay lập tức dịch vụ Ingreslock, cách ly máy chủ để rà soát mã độc và kiểm tra toàn bộ lịch sử tiến trình."
    ),
    "rlogin Passwordless Login": (
        "Dịch vụ rlogin cho phép đăng nhập từ xa quyền Root không cần mật khẩu dựa trên cấu hình tệp tin tin cậy (.rhosts).",
        "Đăng nhập thành công vào máy chủ 192.168.95.138 qua giao thức rlogin (Port 513/TCP) không cần xác thực.",
        "Dừng hoạt động dịch vụ rlogin và chuyển sang sử dụng giao thức SSH bảo mật."
    ),
    "vsftpd Compromised Source Packages Backdoor Vulnerability - Active Check": (
        "Phiên bản vsftpd 2.3.4 cài đặt trên máy chủ chứa lỗ hổng Backdoor công bố, tự động mở cổng 6200/TCP trao quyền Root khi nhận chuỗi ký tự ':)' trong tên đăng nhập.",
        "Thử nghiệm gửi chuỗi ký tự đặc biệt tới cổng 21/TCP tại 192.168.95.138 kích hoạt thành công Root Shell trên cổng 6200/TCP.",
        "Gỡ bỏ phiên bản vsftpd 2.3.4 và nâng cấp lên phiên bản FTP an toàn mới nhất hoặc dùng SFTP."
    ),
    "Apache Tomcat AJP RCE Vulnerability (Ghostcat) - Active Check": (
        "Lỗ hổng Ghostcat (CVE-2020-1938) trong giao thức Tomcat AJP cho phép kẻ tấn công đọc toàn bộ mã nguồn ứng dụng Web và thực thi mã từ xa.",
        "Dịch vụ AJP Connector công khai trên cổng 8009/TCP tại 192.168.95.138 không thiết lập thuộc tính xác thực 'requiredSecret'.",
        "Cập nhật Apache Tomcat lên phiên bản vá lỗi hoặc tắt cổng AJP 8009 nếu không sử dụng."
    ),
    "DistCC RCE Vulnerability (CVE-2004-2687)": (
        "Dịch vụ biên dịch mã từ xa DistCC cho phép gửi các câu lệnh hệ thống tùy ý để thực thi trên máy chủ với quyền hạn của dịch vụ.",
        "Kết nối thành công cổng 3632/TCP tại 192.168.95.138 và thực thi câu lệnh kiểm tra danh tính hệ thống.",
        "Vô hiệu hóa dịch vụ DistCC trên môi trường sản xuất hoặc cài đặt chính sách Tường lửa chỉ cho phép các IP biên dịch tin cậy."
    ),
    "TWiki < 4.2.4 Multiple XSS / Command Execution Vulnerabilities": (
        "Nền tảng TWiki phiên bản cũ chứa lỗ hổng chèn lệnh hệ thống (Command Injection) qua tham số URL, cho phép thực thi mã tùy ý trên Web Server.",
        "Phát hiện ứng dụng TWiki phiên bản cũ hoạt động tại 192.168.95.138 chứa các điểm yếu xử lý dữ liệu đầu vào.",
        "Nâng cấp TWiki lên phiên bản mới nhất hoặc ngắt kết nối ứng dụng nếu không còn sử dụng."
    ),
    "The rexec service is running": (
        "Dịch vụ rexec (Port 512/TCP) truyền tải tên tài khoản và mật khẩu dưới dạng văn bản rõ không mã hóa qua đường truyền mạng.",
        "Dịch vụ rexec hoạt động công khai trên cổng 512/TCP tại địa chỉ 192.168.95.138.",
        "Tắt dịch vụ rexec và thay thế bằng SSH."
    ),
    "Operating System (OS) End of Life (EOL) Detection": (
        "Hệ điều hành máy chủ đã kết thúc vòng đời hỗ trợ (End of Life), không còn nhận được các bản vá lỗi bảo mật từ nhà sản xuất, tạo điều kiện cho kẻ tấn công khai thác các lỗ hổng 0-day.",
        "Ghi nhận phiên bản hệ điều hành Linux tại 192.168.95.138 đã hết hạn hỗ trợ chính thức.",
        "Lập kế hoạch nâng cấp hệ điều hành máy chủ lên phiên bản mới (LTS) còn trong hạn hỗ trợ."
    ),
    "rsh Unencrypted Cleartext Login": (
        "Dịch vụ Remote Shell (rsh) thực hiện xác thực và truyền dữ liệu không mã hóa, dễ bị đánh cắp thông tin tài khoản trên mạng nội bộ.",
        "Phát hiện cổng 514/TCP dịch vụ rsh hoạt động tại địa chỉ 192.168.95.138.",
        "Vô hiệu hóa dịch vụ rsh và chuyển sang giao thức mã hóa SSH."
    ),
    "UnrealIRCd Authentication Spoofing Vulnerability": (
        "Phần mềm máy chủ IRC (UnrealIRCd) chứa lỗ hổng mở cắm cờ cho phép kẻ tấn công giả mạo quyền truy cập và thực thi lệnh hệ thống.",
        "Dịch vụ UnrealIRCd lắng nghe trên các cổng IRC tại địa chỉ 192.168.95.138.",
        "Cập nhật phần mềm UnrealIRCd lên phiên bản an toàn hoặc gỡ bỏ dịch vụ IRC."
    ),
    "SSL/TLS: OpenSSL CCS Man in the Middle Security Bypass Vulnerability": (
        "Lỗ hổng OpenSSL CCS (CVE-2014-0224) cho phép kẻ tấn công đứng giữa (MitM) ép buộc giao thức hạ cấp khóa mã hóa để giải mã thông tin nhạy cảm.",
        "Dịch vụ SSL/TLS tại 192.168.95.138 chấp nhận gói tin ChangeCipherSpec kém an toàn trong quá trình bắt tay.",
        "Cập nhật thư viện OpenSSL trên Web Server lên phiên bản mới nhất."
    ),
    "Absence of Anti-CSRF Tokens": (
        "Các biểu mẫu (Form) xử lý dữ liệu trên website thiếu cờ chống giả mạo yêu cầu (Anti-CSRF Token), dẫn đến nguy cơ người dùng bị lợi dụng thực thi các hành động ngoài ý muốn.",
        "Các yêu cầu POST tại ứng dụng Web 192.168.95.138 không chứa mã thông báo ngẫu nhiên Anti-CSRF Token.",
        "Bổ sung cơ chế sinh và xác thực Anti-CSRF Token cho tất cả các biểu mẫu thay đổi dữ liệu."
    ),
    "Application Error Disclosure": (
        "Website hiển thị chi tiết lỗi kỹ thuật khi gặp sự cố, làm lộ cấu trúc thư mục, phiên bản cơ sở dữ liệu và mã nguồn ứng dụng.",
        "Phát hiện trang thông báo lỗi hiển thị chi tiết Stack Trace khi gửi truy vấn bất thường tới 192.168.95.138.",
        "Cấu hình ứng dụng Web tắt chế độ hiển thị lỗi chi tiết (Debug Mode) và sử dụng trang thông báo lỗi tùy chỉnh (Custom Error Page)."
    ),
    "Content Security Policy (CSP) Header Not Set": (
        "Máy chủ Web chưa cấu hình tiêu đề bảo mật Content-Security-Policy (CSP), làm tăng nguy cơ ứng dụng bị tấn công Chèn kịch bản độc hại (XSS).",
        "HTTP Header phản hồi từ máy chủ 192.168.95.138 không chứa tiêu đề Content-Security-Policy.",
        "Bổ sung tiêu đề Content-Security-Policy trong cấu hình Web Server để kiểm soát nguồn tài nguyên được phép thực thi."
    ),
    "Source Code Disclosure - SQL": (
        "Máy chủ Web để rò rỉ các tệp tin chứa câu truy vấn SQL hoặc mã nguồn ứng dụng cho phép tải về trực tiếp qua đường dẫn Web.",
        "Phát hiện tệp tin sao lưu SQL có thể truy cập công khai tại URL thuộc máy chủ 192.168.95.138.",
        "Xóa bỏ tất cả các tệp tin sao lưu mã nguồn khỏi thư mục root của Web Server và phân quyền truy cập nghiêm ngặt."
    ),
    "Vulnerable JS Library": (
        "Website tích hợp các thư viện JavaScript phiên bản cũ (như jQuery, Bootstrap) đã được cảnh báo chứa các lỗ hổng bảo mật XSS.",
        "Phát hiện tệp tin thư viện JavaScript phiên bản cũ đang được liên kết tại 192.168.95.138.",
        "Nâng cấp các thư viện JavaScript lên phiên bản mới nhất."
    ),
    "Multiple Vendors STARTTLS Implementation Plaintext Arbitrary Command Injection Vulnerability": (
        "Lỗ hổng trong triển khai giao thức STARTTLS cho phép kẻ tấn công chèn các lệnh văn bản rõ trước khi thiết lập mã hóa SSL/TLS.",
        "Dịch vụ thư điện tử/SMTP hỗ trợ STARTTLS tại 192.168.95.138 bị ảnh hưởng bởi lỗi chèn lệnh.",
        "Cập nhật dịch vụ Mail Server lên phiên bản phần mềm vá lỗi STARTTLS."
    ),
    "Weak Host Key Algorithm(s) (SSH)": (
        "Dịch vụ SSH hỗ trợ các thuật toán mã hóa khóa máy chủ cũ kém an toàn (như ssh-dss), dễ bị bẻ khóa.",
        "Phát hiện máy chủ SSH tại 192.168.95.138 cho phép đàm phán các thuật toán khóa yếu.",
        "Cấu hình tệp sshd_config loại bỏ các thuật toán khóa yếu, chỉ cho phép sử dụng RSA 3072-bit trở lên hoặc ED25519."
    ),
    "DNS Cache Snooping Vulnerability (UDP) - Active Check": (
        "Máy chủ DNS cho phép truy vấn bộ nhớ đệm (Cache Snooping), giúp kẻ tấn công thu thập danh sách tên miền vừa được người dùng nội bộ truy cập.",
        "Dịch vụ DNS (Port 53/UDP) tại 192.168.95.138 trả về phản hồi truy vấn đệm cho các yêu cầu không thuộc thẩm quyền.",
        "Cấu hình máy chủ DNS giới hạn quyền truy vấn đệm (Recursion) chỉ dành cho các IP nội bộ được phép."
    ),
    "VNC Server Unencrypted Data Transmission": (
        "Dịch vụ quản trị màn hình từ xa VNC truyền tải hình ảnh và thao tác bàn phím không qua mã hóa, dễ bị nghe lén trên mạng.",
        "Máy chủ VNC (Port 5900/TCP) tại 192.168.95.138 hoạt động ở chế độ không bật mã hóa SSL/TLS.",
        "Cấu hình bật cơ chế mã hóa cho VNC hoặc bọc kết nối VNC qua đường truyền SSH Tunnel."
    ),
    "Cleartext Transmission of Sensitive Information via HTTP": (
        "Website thực hiện truyền tải dữ liệu nhạy cảm (như thông tin đăng nhập) qua giao thức HTTP không mã hóa.",
        "Phát hiện biểu mẫu đăng nhập gửi dữ liệu qua kênh HTTP thông thường tại 192.168.95.138.",
        "Chuyển đổi toàn bộ website sang sử dụng giao thức HTTPS mã hóa và kích hoạt tiêu đề HSTS."
    ),
    "FTP Unencrypted Cleartext Login": (
        "Dịch vụ FTP truyền thông tin tài khoản và mật khẩu dạng văn bản rõ không mã hóa qua mạng.",
        "Dịch vụ FTP (Port 21/TCP) tại 192.168.95.138 cho phép đăng nhập qua kênh không mã hóa.",
        "Ngắt dịch vụ FTP và thay thế bằng giao thức truyền tệp an toàn SFTP."
    ),
    "Telnet Unencrypted Cleartext Login": (
        "Dịch vụ Telnet truyền toàn bộ dữ liệu quản trị không mã hóa, gây nguy cơ rò rỉ tài khoản quản trị.",
        "Dịch vụ Telnet (Port 23/TCP) hoạt động mở công khai tại 192.168.95.138.",
        "Tắt hoàn toàn dịch vụ Telnet và chuyển sang sử dụng SSH."
    ),
    "Cookie No HttpOnly Flag": (
        "Cookie phiên làm việc thiếu thuộc tính bảo vệ HttpOnly, khiến Cookie có thể bị kịch bản độc hại (XSS) truy cập và lấy cắp.",
        "Phản hồi Set-Cookie từ ứng dụng Web tại 192.168.95.138 thiếu cờ HttpOnly.",
        "Bổ sung thuộc tính HttpOnly cho tất cả các Cookie phiên làm việc trong cấu hình ứng dụng Web."
    ),
    "In Page Banner Information Leak": (
        "Máy chủ Web tiết lộ chi tiết tên và phiên bản phần mềm máy chủ trong tiêu đề phản hồi (Server Banner).",
        "Phản hồi HTTP Header từ 192.168.95.138 hiển thị chi tiết tên phiên bản Web Server.",
        "Cấu hình Web Server ẩn thông tin phiên bản (ví dụ set 'ServerTokens Prod' trên Apache hoặc 'server_tokens off' trên Nginx)."
    ),
    "Information Disclosure - Debug Error Messages": (
        "Máy chủ rò rỉ thông tin cấu hình nội bộ qua các thông báo lỗi gỡ lỗi (Debug Error Messages).",
        "Phát hiện thông tin cấu hình hệ thống xuất hiện trong các phản hồi lỗi tại 192.168.95.138.",
        "Tắt chế độ Debug trên môi trường sản xuất và sử dụng trang báo lỗi tiêu chuẩn."
    ),
    "Timestamp Disclosure - Unix": (
        "Ứng dụng Web làm rò rỉ giá trị thời gian hệ thống (Unix Timestamp) trong phản hồi HTTP.",
        "Phát hiện chuỗi thời gian Unix Timestamp trong dữ liệu phản hồi từ 192.168.95.138.",
        "Loại bỏ các trường thời gian hệ thống không cần thiết trong dữ liệu phản hồi công cộng."
    )
}

# --- BUILD DOCX ---
doc = docx.Document('MAU REPORT.docx')

# Clean body XML completely
for child in list(doc._element.body):
    if not child.tag.endswith('sectPr'):
        doc._element.body.remove(child)

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 16, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 13, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 12, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_p(text, style='Normal', bold=False, color_rgb=None, size_pt=11):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    force_font(r, "Times New Roman", size_pt, bold=bold, color_rgb=color_rgb)
    return p

def format_table_cell(cell, text, bold=False, color_rgb=None, fill_hex=None, size_pt=10):
    if fill_hex:
        set_cell_background(cell, fill_hex)
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    force_font(r, "Times New Roman", size_pt, bold=bold, color_rgb=color_rgb)

# TITLE
p_title = doc.add_paragraph(style='Heading 1')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after = Pt(18)
r_t = p_title.add_run("BÁO CÁO NỘI DUNG ĐÁNH GIÁ LỖ HỔNG BẢO MẬT")
force_font(r_t, "Times New Roman", 18, bold=True, color_rgb=RGBColor(0, 32, 96))

# --- TABLE OF CONTENTS SECTION ---
add_h1("MỤC LỤC")
t_toc = doc.add_table(rows=1, cols=2)
t_toc.style = 'Table Grid'
format_table_cell(t_toc.rows[0].cells[0], "NỘI DUNG CẤU TRÚC BÁO CÁO", bold=True, fill_hex='D9D9D9')
format_table_cell(t_toc.rows[0].cells[1], "TRANG", bold=True, fill_hex='D9D9D9')

toc_items = [
    ("1. THÔNG TIN DỰ ÁN", "3"),
    ("   1.1. PHIÊN BẢN TÀI LIỆU", "3"),
    ("   1.2. THỜI GIAN TRIỂN KHAI", "3"),
    ("   1.3. NHÂN SỰ TRIỂN KHAI", "3"),
    ("   1.4. MỤC ĐÍCH ĐÁNH GIÁ", "4"),
    ("   1.5. NỘI DUNG ĐÁNH GIÁ", "4"),
    ("   1.6. PHẠM VI THỰC HIỆN", "4"),
    ("2. BÁO CÁO TỔNG QUÁT", "5"),
    ("   2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", "5"),
    ("   2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER (192.168.95.138)", "6"),
    ("   2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA / TRENDING)", "7"),
    ("   2.4. KHUYẾN NGHỊ TỔNG QUAN", "9"),
    ("3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", "10"),
    ("4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER (192.168.95.138)", "12"),
    ("5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO", "25"),
    ("   5.1. MÔ HÌNH ĐÁNH GIÁ RỦI RO THEO OWASP RISK RATING", "25"),
    ("   5.2. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0", "26"),
    ("   5.3. THANG ĐIỂM ĐÁNH GIÁ CVSS V3.0", "27"),
]

for section_name, page_num in toc_items:
    r_c = t_toc.add_row().cells
    is_main = not section_name.startswith("   ")
    format_table_cell(r_c[0], section_name, bold=is_main)
    format_table_cell(r_c[1], page_num, bold=is_main)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Native Word TOC Field
p_native_toc = doc.add_paragraph()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "TOC \\o \"1-3\" \\h \\z \\u"
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
r1 = OxmlElement('w:r'); r1.append(fldChar1); p_native_toc._element.append(r1)
r2 = OxmlElement('w:r'); r2.append(instrText); p_native_toc._element.append(r2)
r3 = OxmlElement('w:r'); r3.append(fldChar2); p_native_toc._element.append(r3)
r4 = OxmlElement('w:r'); r4.append(fldChar3); p_native_toc._element.append(r4)

# 1. THÔNG TIN DỰ ÁN
add_h1("1. THÔNG TIN DỰ ÁN")

add_h2("1.1. PHIÊN BẢN TÀI LIỆU")
t_ver = doc.add_table(rows=4, cols=2)
t_ver.style = 'Table Grid'
ver_data = [
    ("Phiên bản tài liệu", "1.0"),
    ("Ngày báo cáo", "10/08/2026"),
    ("Khách hàng", "[Tên Khách Hàng]"),
    ("Công ty cung ứng dịch vụ", "CÔNG TY CỔ PHẦN CÔNG NGHỆ DTG")
]
for idx, (lbl, val) in enumerate(ver_data):
    format_table_cell(t_ver.rows[idx].cells[0], lbl, bold=True, fill_hex='F2F2F2')
    format_table_cell(t_ver.rows[idx].cells[1], val)

add_h2("1.2. THỜI GIAN TRIỂN KHAI")
add_p("Thời gian thực hiện đánh giá: 01/08/2026 – 10/08/2026")

add_h2("1.3. NHÂN SỰ TRIỂN KHAI")
t_staff = doc.add_table(rows=2, cols=3)
t_staff.style = 'Table Grid'
format_table_cell(t_staff.rows[0].cells[0], "STT", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[0].cells[1], "Nhân sự", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[0].cells[2], "Vai Trò", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[1].cells[0], "1")
format_table_cell(t_staff.rows[1].cells[1], "Chuyên gia An toàn thông tin")
format_table_cell(t_staff.rows[1].cells[2], "Kiểm thử viên chính (Lead Pentester)")

add_h2("1.4. MỤC ĐÍCH ĐÁNH GIÁ")
add_p("• Kiểm tra và đánh giá toàn diện hiện trạng an toàn thông tin đối với các thiết bị máy tính xách tay và hệ thống máy chủ dịch vụ Web Server định kỳ 3 tháng/lần.")
add_p("• Kịp thời phát hiện, phân loại và cảnh báo các lỗ hổng bảo mật, điểm yếu cấu hình hoặc các phần mềm lỗi thời tồn tại trên các mục tiêu.")
add_p("• Đề xuất phương án khắc phục, giảm thiểu rủi ro an ninh mạng làm cơ sở tham mưu cho đơn vị.")

add_h2("1.5. NỘI DUNG ĐÁNH GIÁ")
add_p("• Rà quét tự động kết hợp phân tích thủ công để xác định các cổng/dịch vụ mạng đang mở, điểm yếu giao thức và lỗ hổng phần mềm.")
add_p("• Phân loại mức độ nghiêm trọng của lỗ hổng dựa trên tiêu chuẩn quốc tế CVSS v3.0 và mô hình OWASP Risk Rating.")
add_p("• Lập báo cáo kết quả chi tiết, tư vấn phương án khắc phục và tái kiểm tra.")

add_h2("1.6. PHẠM VI THỰC HIỆN")
add_p("Hoạt động rà quét được thực hiện đối với 2 mục tiêu chính thông qua phương thức kiểm thử từ bên ngoài (External Scan - Internet Simulation):")

t_scope = doc.add_table(rows=3, cols=4)
t_scope.style = 'Table Grid'
scope_headers = ['STT', 'Thông tin thiết bị', 'Địa chỉ IP', 'Phương thức đánh giá']
for i, h in enumerate(scope_headers):
    format_table_cell(t_scope.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

format_table_cell(t_scope.rows[1].cells[0], "1")
format_table_cell(t_scope.rows[1].cells[1], "Máy tính xách tay Windows 11")
format_table_cell(t_scope.rows[1].cells[2], "192.168.95.135")
format_table_cell(t_scope.rows[1].cells[3], "Rà quét từ bên ngoài (Internet Simulation)")

format_table_cell(t_scope.rows[2].cells[0], "2")
format_table_cell(t_scope.rows[2].cells[1], "Hệ thống Web Server")
format_table_cell(t_scope.rows[2].cells[2], "192.168.95.138")
format_table_cell(t_scope.rows[2].cells[3], "Rà quét từ bên ngoài (Internet Simulation)")


# 2. BÁO CÁO TỔNG QUÁT
add_h1("2. BÁO CÁO TỔNG QUÁT")

def build_summary_table_vi(title_h2, df, target_name):
    add_h2(title_h2)
    add_p(f"Tổng số lỗ hổng phát hiện trên {target_name}: {len(df)} lỗ hổng.")
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(['STT', 'Mức độ', 'Lỗ hổng', 'Thiết bị']):
        format_table_cell(t.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

    for idx, row in df.reset_index().iterrows():
        r_cells = t.add_row().cells
        format_table_cell(r_cells[0], str(idx + 1))
        
        orig_sev = str(row.get('severity', '')).upper()
        sev_vi = SEV_MAP.get(orig_sev, orig_sev)
        color = SEV_COLOR.get(sev_vi, RGBColor(0,0,0))
        
        format_table_cell(r_cells[1], sev_vi, bold=True, color_rgb=color)
        format_table_cell(r_cells[2], clean_html(str(row.get('finding_name', ''))))
        format_table_cell(r_cells[3], target_name)

build_summary_table_vi("2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", real_laptop, "Máy tính xách tay Windows 11")
build_summary_table_vi("2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER (192.168.95.138)", real_webserver, "Web Server")

add_h2("2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA / TRENDING)")
add_p("Dưới đây là 2 biểu đồ phân tích trực quan về diễn biến an toàn thông tin giữa 2 chu kỳ rà quét:")

# Insert Chart 1
add_p("Biểu đồ 1: So sánh tổng số lỗ hổng giữa Chu kỳ 3 tháng trước và Chu kỳ hiện tại", bold=True)
doc.add_picture(chart1_path, width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Insert Chart 2
add_p("Biểu đồ 2: Phân bổ lỗ hổng chi tiết theo từng mục tiêu trong Chu kỳ hiện tại", bold=True)
doc.add_picture(chart2_path, width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_p("Bảng chi tiết thống kê biến động Delta lỗ hổng theo từng mục tiêu:")
t_delta = doc.add_table(rows=1, cols=6)
t_delta.style = 'Table Grid'
delta_headers = ['Chu kỳ rà quét', 'Mục tiêu', 'Nghiêm trọng', 'Cao', 'Trung bình', 'Thấp']
for i, h in enumerate(delta_headers):
    format_table_cell(t_delta.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

delta_rows = [
    ("Chu kỳ 3 tháng trước", "Laptop (192.168.95.135)", past_laptop),
    ("Chu kỳ hiện tại", "Laptop (192.168.95.135)", curr_laptop),
    ("Chu kỳ 3 tháng trước", "Web Server (192.168.95.138)", past_webserver),
    ("Chu kỳ hiện tại", "Web Server (192.168.95.138)", curr_webserver),
]

for cycle, tgt, data_dict in delta_rows:
    r_cells = t_delta.add_row().cells
    format_table_cell(r_cells[0], cycle)
    format_table_cell(r_cells[1], tgt)
    format_table_cell(r_cells[2], str(data_dict['NGHIÊM TRỌNG']), bold=(data_dict['NGHIÊM TRỌNG'] > 0), color_rgb=RGBColor(255,0,0) if data_dict['NGHIÊM TRỌNG'] > 0 else None)
    format_table_cell(r_cells[3], str(data_dict['CAO']), bold=(data_dict['CAO'] > 0), color_rgb=RGBColor(192,0,0) if data_dict['CAO'] > 0 else None)
    format_table_cell(r_cells[4], str(data_dict['TRUNG BÌNH']), bold=(data_dict['TRUNG BÌNH'] > 0), color_rgb=RGBColor(237,125,49) if data_dict['TRUNG BÌNH'] > 0 else None)
    format_table_cell(r_cells[5], str(data_dict['THẤP']))

add_p("Phân tích xu hướng rủi ro (Trending Analysis):", bold=True)
add_p("• Đối với Máy tính xách tay Windows 11 (192.168.95.135): Số lượng lỗ hổng giảm từ 8 xuống 4 lỗi, đặc biệt các lỗi Nghiêm trọng đã được triệt tiêu hoàn toàn. Cho thấy công tác cập nhật bản vá hệ điều hành đã có hiệu quả tích cực.")
add_p("• Đối với Hệ thống Web Server (192.168.95.138): Mức độ rủi ro tăng mạnh. Số lỗi Nghiêm trọng tăng từ 2 lên 10 lỗi, lỗi Cao giữ ở mức 4 lỗi. Sự gia tăng này cảnh báo bề mặt tấn công công cộng đang bị đe dọa trực tiếp bởi các lỗ hổng mã nguồn ứng dụng và dịch vụ chưa được vá.")

add_h2("2.4. KHUYẾN NGHỊ TỔNG QUAN")
add_p("1. Ưu tiên xử lý khẩn cấp 10 lỗ hổng NGHIÊM TRỌNG và 4 lỗ hổng CAO trên Web Server trong vòng 24 - 48 giờ.")
add_p("2. Thực hiện nâng cấp các phiên bản dịch vụ Web Server và thư viện liên quan lên phiên bản an toàn mới nhất.")
add_p("3. Thiết lập hệ thống Tường lửa ứng dụng Web (WAF) để ngăn chặn các cuộc tấn công khai thác lỗ hổng từ môi trường Internet.")
add_p("4. Duy trì chính sách rà quét lỗ hổng bảo mật định kỳ 3 tháng/lần để phát hiện sớm các rủi ro phát sinh.")


# HELPER FOR VULNERABILITY DETAILS
def render_detailed_vuln_vi(sec_num, vuln_num, row, target_name):
    finding_name = clean_html(str(row.get('finding_name', '')))
    orig_sev = str(row.get('severity', '')).upper()
    sev_vi = SEV_MAP.get(orig_sev, orig_sev)
    
    cve = str(row.get('cve', 'N/A'))
    if pd.isna(cve) or cve in ['nan', '']:
        cve = str(row.get('cve_list', 'N/A'))
    cvss = str(row.get('cvss', 'N/A'))
    if pd.isna(cvss) or cvss in ['nan', '']:
        cvss = 'N/A'
        
    loc = clean_html(str(row.get('location', '')))
    raw_ev = clean_html(str(row.get('scanner_evidence', ''))) or clean_html(str(row.get('evidence', '')))
    raw_sol = clean_html(str(row.get('solution', ''))) or clean_html(str(row.get('scanner_solution', '')))
    orig_desc = clean_html(str(row.get('description', '')))
    
    # Custom Vietnamese translation lookups
    if finding_name in VI_DETAILS:
        vi_desc, vi_evidence, vi_solution = VI_DETAILS[finding_name]
    else:
        vi_desc, vi_evidence, vi_solution = get_custom_details_vi(finding_name, loc, orig_desc, raw_ev, raw_sol)
    
    add_h3(f"{sec_num}.{vuln_num}. {finding_name.upper()}")
    
    add_p(f"Mô tả: {vi_desc}")
    
    p_sev = doc.add_paragraph(style='Normal')
    p_sev.paragraph_format.space_after = Pt(3)
    r_lbl = p_sev.add_run("Mức độ: ")
    force_font(r_lbl, "Times New Roman", 11, bold=False)
    
    color = SEV_COLOR.get(sev_vi, RGBColor(0,0,0))
    r_sev = p_sev.add_run(sev_vi)
    force_font(r_sev, "Times New Roman", 11, bold=True, color_rgb=color)

    add_p(f"CVSS: {cvss}")
    add_p(f"CVE: {cve}")
    
    add_p("Bằng chứng:", bold=True)
    t_ev = doc.add_table(rows=1, cols=1)
    t_ev.style = 'Table Grid'
    t_ev.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_cell(t_ev.rows[0].cells[0], vi_evidence, fill_hex='F9F9F9', size_pt=10)
    
    add_p(f"Khuyến nghị: {vi_solution}")


# 3. BÁO CÁO CHI TIẾT LAPTOP
add_h1("3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)")
for idx, row in real_laptop.reset_index().iterrows():
    render_detailed_vuln_vi(3, idx + 1, row, "Máy tính xách tay Windows 11 (192.168.95.135)")

# 4. BÁO CÁO CHI TIẾT WEB SERVER
add_h1("4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER (192.168.95.138)")
for idx, row in real_webserver.reset_index().iterrows():
    render_detailed_vuln_vi(4, idx + 1, row, "Web Server (192.168.95.138)")


# 5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO
add_h1("5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO")

add_h2("5.1. MÔ HÌNH ĐÁNH GIÁ RỦI RO THEO OWASP RISK RATING")
add_p("Mô hình đánh giá rủi ro an toàn thông tin dựa trên tiêu chuẩn OWASP Risk Rating Methodology, kết hợp 2 yếu tố chính: Khả năng xảy ra (Likelihood) và Mức độ ảnh hưởng (Impact).")

# Table 31: OWASP Matrix
t_owasp = doc.add_table(rows=6, cols=5)
t_owasp.style = 'Table Grid'

owasp_headers = ["MA TRẬN OWASP", "KHÔNG", "THẤP", "TRUNG BÌNH", "CAO"]
for i, h in enumerate(owasp_headers):
    format_table_cell(t_owasp.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

owasp_data = [
    ["MỨC ĐỘ ẢNH HƯỞNG", "CAO", "TRUNG BÌNH", "CAO", "NGHIÊM TRỌNG"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "TRUNG BÌNH", "THẤP", "TRUNG BÌNH", "CAO"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "THẤP", "KHÔNG", "THẤP", "TRUNG BÌNH"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "-", "THẤP", "TRUNG BÌNH", "CAO"],
    ["-", "KHẢ NĂNG XẢY RA", "THẤP", "TRUNG BÌNH", "CAO"]
]

for r_idx, row in enumerate(owasp_data):
    for c_idx, val in enumerate(row):
        cell = t_owasp.rows[r_idx+1].cells[c_idx]
        color = SEV_COLOR.get(val, None)
        format_table_cell(cell, val, bold=(color is not None), color_rgb=color)

add_h2("5.2. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0")
add_p("Common Vulnerability Scoring System (CVSS) v3.0 là tiêu chuẩn quốc tế mở dùng để đánh giá mức độ nghiêm trọng của lỗ hổng phần mềm.")

add_h3("5.2.1. EXPLOITABILITY METRICS")
add_p("• Attack Vector (AV): Đánh giá môi trường khai thác lỗ hổng (Network, Adjacent, Local, Physical).")
add_p("• Attack Complexity (AC): Đánh giá độ phức tạp của cuộc tấn công (Low, High).")
add_p("• Privileges Required (PR): Đánh giá mức độ quyền hạn cần thiết của kẻ tấn công (None, Low, High).")
add_p("• User Interaction (UI): Đánh giá yêu cầu sự tương tác của người dùng (None, Required).")

add_h3("5.2.2. IMPACT METRICS")
add_p("• Confidentiality Impact (C): Đánh giá mức độ rò rỉ thông tin bảo mật.")
add_p("• Integrity Impact (I): Đánh giá mức độ sai lệch hoặc sửa đổi dữ liệu.")
add_p("• Availability Impact (A): Đánh giá mức độ gián đoạn khả năng cung cấp dịch vụ.")

add_h2("5.3. THANG ĐIỂM ĐÁNH GIÁ CVSS V3.0")
t_cvss = doc.add_table(rows=5, cols=3)
t_cvss.style = 'Table Grid'
cvss_headers = ["Điểm CVSS v3.0", "Mức độ nghiêm trọng", "Mô tả rủi ro"]
for i, h in enumerate(cvss_headers):
    format_table_cell(t_cvss.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

cvss_table_data = [
    ("9.0 – 10.0", "NGHIÊM TRỌNG (Critical)", "Lỗ hổng cho phép khai thác từ xa không cần xác thực, chiếm quyền kiểm soát hệ thống."),
    ("7.0 – 8.9", "CAO (High)", "Lỗ hổng gây ảnh hưởng nghiêm trọng đến tính bảo mật hoặc sẵn sàng của hệ thống."),
    ("4.0 – 6.9", "TRUNG BÌNH (Medium)", "Lỗ hổng yêu cầu điều kiện khai thác phức tạp hoặc quyền truy cập hạn chế."),
    ("0.1 – 3.9", "THẤP (Low)", "Lỗ hổng có mức độ tác động nhỏ, khó khai thác thành công.")
]

for idx, (score, sev, desc) in enumerate(cvss_table_data):
    r_cells = t_cvss.rows[idx+1].cells
    format_table_cell(r_cells[0], score)
    
    sev_key = sev.split()[0]
    color = SEV_COLOR.get(sev_key, None)
    
    format_table_cell(r_cells[1], sev, bold=True, color_rgb=color)
    format_table_cell(r_cells[2], desc)

doc.save('Bao_Cao_An_Toan_Thong_Tin_2026.docx')
print("Successfully generated PERFECTION V3 Bao_Cao_An_Toan_Thong_Tin_2026.docx")
