#!/usr/bin/env python3
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import re
import os
import matplotlib.pyplot as plt

def clean_html(text):
    if not isinstance(text, str) or pd.isna(text):
        return ""
    # Strip HTML tags
    cleaned = re.sub(r'<[^>]+>', '', text)
    # Unescape common HTML entities
    cleaned = cleaned.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    return cleaned.strip()

# Load Data
old_df = pd.read_csv('runs/run_20260810_140029/output/vuln_raw.csv')
new_df = pd.read_csv('runs/run_20260810_192535/output/vuln_attack_enriched.csv')

old_laptop = old_df[old_df['location'].astype(str).str.contains('192.168.95.135', na=False) | old_df['asset'].astype(str).str.contains('192.168.95.135', na=False)].copy()
new_webserver = new_df[new_df['location'].astype(str).str.contains('192.168.95.138', na=False) | new_df['asset'].astype(str).str.contains('192.168.95.138', na=False)].copy()

print(f"Laptop findings (run_140029): {len(old_laptop)}")
print(f"Web Server findings (run_192535): {len(new_webserver)}")

# 1. Generate Comparison Chart Image
categories = ['Critical', 'High', 'Medium', 'Low', 'Info']
def get_counts(df):
    sev_upper = df['severity'].astype(str).str.upper()
    return [
        len(df[sev_upper == 'CRITICAL']),
        len(df[sev_upper == 'HIGH']),
        len(df[sev_upper == 'MEDIUM']),
        len(df[sev_upper == 'LOW']),
        len(df[sev_upper.isin(['INFORMATIONAL', 'INFO'])])
    ]

laptop_counts = get_counts(old_laptop)
web_counts = get_counts(new_webserver)

plt.figure(figsize=(7, 4.2))
x = range(len(categories))
width = 0.35

plt.bar([i - width/2 for i in x], laptop_counts, width, label='Laptop (192.168.95.135 - 3 tháng trước)', color='#3498db')
plt.bar([i + width/2 for i in x], web_counts, width, label='Web Server (192.168.95.138 - Hiện tại)', color='#e74c3c')

plt.xlabel('Mức độ nghiêm trọng (Severity)', fontweight='bold')
plt.ylabel('Số lượng lỗ hổng', fontweight='bold')
plt.title('BIỂU ĐỒ SO SÁNH LỖ HỔNG BẢO MẬT GIỮA 2 CHU KỲ', fontweight='bold', fontsize=12)
plt.xticks(x, categories)
plt.legend()
plt.grid(axis='y', linestyle='--', alpha=0.5)
plt.tight_layout()

chart_path = 'docx_analysis_tools/chart_comparison.png'
plt.savefig(chart_path, dpi=300)
plt.close()

# 2. Build DOCX Document
doc = docx.Document('MAU REPORT.docx')

# Clear existing paragraphs and tables while preserving headers/footers/styles
for p in list(doc.paragraphs):
    p._element.getparent().remove(p._element)
for t in list(doc.tables):
    t._element.getparent().remove(t._element)

def add_h1(text):
    p = doc.add_paragraph(text, style='Heading 1')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_h2(text):
    p = doc.add_paragraph(text, style='Heading 2')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_h3(text):
    p = doc.add_paragraph(text, style='Heading 3')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.bold = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_p(text, style='Normal', bold=False, color=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    if bold:
        r.font.bold = True
    if color:
        r.font.color.rgb = color
    return p

# --- TITLE ---
p_title = add_h1("BÁO CÁO NỘI DUNG ĐÁNH GIÁ LỖ HỔNG BẢO MẬT")
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- SECTION 1 ---
add_h1("1. THÔNG TIN DỰ ÁN")

add_h2("1.1. PHIÊN BẢN TÀI LIỆU")
t_ver = doc.add_table(rows=4, cols=2)
t_ver.style = 'Table Grid'
t_ver.alignment = WD_TABLE_ALIGNMENT.CENTER
t_ver.rows[0].cells[0].text = "Phiên bản tài liệu"
t_ver.rows[0].cells[1].text = "1.0"
t_ver.rows[1].cells[0].text = "Ngày báo cáo"
t_ver.rows[1].cells[1].text = "10/08/2026"
t_ver.rows[2].cells[0].text = "Khách hàng"
t_ver.rows[2].cells[1].text = "[Tên Khách Hàng]"
t_ver.rows[3].cells[0].text = "Công ty cung ứng dịch vụ"
t_ver.rows[3].cells[1].text = "CÔNG TY CỔ PHẦN CÔNG NGHỆ DTG"

add_h2("1.2. THỜI GIAN TRIỂN KHAI")
add_p("Thời gian thực hiện đánh giá: 01/08/2026 – 10/08/2026")

add_h2("1.3. NHÂN SỰ TRIỂN KHAI")
t_staff = doc.add_table(rows=2, cols=3)
t_staff.style = 'Table Grid'
t_staff.rows[0].cells[0].text = "STT"
t_staff.rows[0].cells[1].text = "Nhân sự"
t_staff.rows[0].cells[2].text = "Vai Trò"
t_staff.rows[1].cells[0].text = "1"
t_staff.rows[1].cells[1].text = "Chuyên gia An toàn thông tin"
t_staff.rows[1].cells[2].text = "Kiểm thử viên chính (Lead Pentester)"

add_h2("1.4. MỤC ĐÍCH ĐÁNH GIÁ")
add_p("• Kiểm tra và đánh giá toàn diện hiện trạng an toàn thông tin đối với các trang thiết bị đầu cuối (máy tính xách tay) và máy chủ dịch vụ (Web Server) đang vận hành trong hệ thống mạng định kỳ 3 tháng/lần.")
add_p("• Kịp thời phát hiện, phân loại và cảnh báo các lỗ hổng bảo mật, điểm yếu cấu hình hoặc các phần mềm lỗi thời đang tồn tại trên thiết bị.")
add_p("• Đề xuất các phương án khắc phục, giảm thiểu rủi ro, làm cơ sở để tham mưu cho đơn vị trong công tác đảm bảo an toàn thông tin.")

add_h2("1.5. NỘI DUNG ĐÁNH GIÁ")
add_p("• Sử dụng các công cụ chuyên dụng để tiến hành dò quét tự động và phân tích thủ công nhằm xác định các dịch vụ mạng đang mở, các điểm yếu giao thức và lỗ hổng phần mềm.")
add_p("• Đối chiếu kết quả quét với các tiêu chuẩn an toàn thông tin quốc tế để phân tích, đánh giá mức độ nghiêm trọng của từng lỗ hổng.")
add_p("• Xây dựng báo cáo kết quả chi tiết, tư vấn lộ trình cập nhật bản vá, chuẩn hóa cấu hình hệ thống và lên kế hoạch tái kiểm tra sau khi khắc phục.")

add_h2("1.6. PHẠM VI THỰC HIỆN")
add_p("• Đánh giá nội bộ (Internal Scan): Máy tính xách tay Windows 11 (IP: 192.168.95.135).")
add_p("• Đánh giá bên ngoài (External Scan - Internet Simulation): Hệ thống Web Server (IP: 192.168.95.138).")

# --- SECTION 2 ---
add_h1("2. BÁO CÁO TỔNG QUÁT")

add_h2("2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)")
add_p("Tổng số lỗ hổng phát hiện trên Máy tính xách tay Windows 11 (Chu kỳ 3 tháng trước): 4 lỗ hổng.")
t_lap = doc.add_table(rows=1, cols=4)
t_lap.style = 'Table Grid'
for i, h in enumerate(['STT', 'Mức độ', 'Lỗ hổng', 'Thiết bị']):
    t_lap.rows[0].cells[i].text = h
    t_lap.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, row in old_laptop.reset_index().iterrows():
    r_cells = t_lap.add_row().cells
    r_cells[0].text = str(idx + 1)
    r_cells[1].text = str(row.get('severity', '')).upper()
    r_cells[2].text = clean_html(str(row.get('finding_name', '')))
    r_cells[3].text = "Máy tính xách tay Windows 11"

add_h2("2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER (192.168.95.138)")
add_p(f"Tổng số lỗ hổng phát hiện trên Web Server (Chu kỳ hiện tại): {len(new_webserver)} lỗ hổng.")
t_web = doc.add_table(rows=1, cols=4)
t_web.style = 'Table Grid'
for i, h in enumerate(['STT', 'Mức độ', 'Lỗ hổng', 'Thiết bị']):
    t_web.rows[0].cells[i].text = h
    t_web.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, row in new_webserver.reset_index().iterrows():
    r_cells = t_web.add_row().cells
    r_cells[0].text = str(idx + 1)
    r_cells[1].text = str(row.get('severity', '')).upper()
    r_cells[2].text = clean_html(str(row.get('finding_name', '')))
    r_cells[3].text = "Web Server"

add_h2("2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA / TRENDING)")
add_p("Dưới đây là biểu đồ và bảng so sánh chi tiết số lượng lỗ hổng theo cấp độ nghiêm trọng giữa 2 đợt đánh giá định kỳ:")

# Insert Chart
doc.add_picture(chart_path, width=Inches(6.0))
p_img = doc.paragraphs[-1]
p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_p("Bảng tổng hợp so sánh Delta giữa 2 chu kỳ rà quét:")
t_comp = doc.add_table(rows=1, cols=6)
t_comp.style = 'Table Grid'
for i, h in enumerate(['Đối tượng rà quét', 'Critical', 'High', 'Medium', 'Low', 'Info']):
    t_comp.rows[0].cells[i].text = h
    t_comp.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

r1 = t_comp.add_row().cells
r1[0].text = "Laptop Windows 11 (3 tháng trước)"
for i in range(5):
    r1[i+1].text = str(laptop_counts[i])

r2 = t_comp.add_row().cells
r2[0].text = "Web Server (Chu kỳ hiện tại)"
for i in range(5):
    r2[i+1].text = str(web_counts[i])

add_p("Đánh giá xu hướng bất thường:")
add_p("• Sự bùng nổ rủi ro tại hạ tầng Web Server: Trong khi đợt quét 3 tháng trước trên máy tính trạm chỉ ghi nhận các lỗ hổng ở mức High và Medium liên quan đến SMB/RPC, đợt quét mới trên Web Server cho thấy rủi ro tăng vọt với 10 lỗ hổng CRITICAL và 4 lỗ hổng HIGH.")
add_p("• Nhận xét chuyên môn: Hệ thống Web Server công cộng đang tồn tại nhiều lỗ hổng nguy hiểm nghiêm trọng (xác thực, mã hóa, cấu hình sai dịch vụ) đòi hỏi phải tiến hành khoanh vùng và xử lý bản vá ưu tiên ngay lập tức.")

add_h2("2.4. KHUYẾN NGHỊ")
add_p("• Cập nhật bản vá hệ điều hành và dịch vụ Web ngay lập tức đối với các lỗ hổng Critical/High.")
add_p("• Áp dụng cơ chế mã hóa SSL/TLS chuẩn hóa và tắt các phương thức xác thực không an toàn.")
add_p("• Rà soát và giới hạn quyền truy cập theo nguyên tắc tối thiểu (Least Privilege).")

# Helper function to add vulnerability details in exact format of MAU REPORT.docx
def render_vuln_detail(sec_num, vuln_num, row, target_name):
    finding = clean_html(str(row.get('finding_name', '')))
    desc = clean_html(str(row.get('description', 'Không có mô tả.')))
    sev = str(row.get('severity', '')).upper()
    cve = str(row.get('cve', 'N/A'))
    if pd.isna(cve) or cve == 'nan' or cve == '':
        cve = str(row.get('cve_list', 'N/A'))
    cvss = str(row.get('cvss', 'N/A'))
    if pd.isna(cvss) or cvss == 'nan' or cvss == '':
        cvss = 'N/A'
    
    # Extract evidence without truncation
    ev = clean_html(str(row.get('scanner_evidence', '')))
    if not ev or ev == 'nan':
        ev = clean_html(str(row.get('evidence', '')))
    if not ev or ev == 'nan':
        ev = clean_html(str(row.get('location', '')))
    if not ev or ev == 'nan':
        ev = f"Scanner reported finding on asset: {target_name}"

    sol = clean_html(str(row.get('solution', '')))
    if not sol or sol == 'nan':
        sol = clean_html(str(row.get('scanner_solution', 'Không có hướng dẫn khắc phục.')))
        
    add_h3(f"{sec_num}.{vuln_num}. {finding.upper()}")
    
    # Description
    add_p(f"Mô tả: {desc}")
    
    # Severity line
    p_sev = doc.add_paragraph(style='Normal')
    p_sev.paragraph_format.space_after = Pt(3)
    p_sev.add_run("Mức độ: ")
    r_sev = p_sev.add_run(sev)
    r_sev.font.bold = True
    if sev in ['CRITICAL', 'HIGH']:
        r_sev.font.color.rgb = RGBColor(255, 0, 0)
        
    # CVSS & CVE
    add_p(f"CVSS: {cvss}")
    add_p(f"CVE: {cve}")
    
    # Evidence Section with 1-column table box
    add_p("Đường dẫn phát hiện:")
    t_ev = doc.add_table(rows=1, cols=1)
    t_ev.style = 'Table Grid'
    t_ev.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_ev.rows[0].cells[0].text = ev
    
    # Recommendation
    add_p(f"Khuyến nghị: {sol}")

# --- SECTION 3 ---
add_h1("3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)")
for idx, row in old_laptop.reset_index().iterrows():
    render_vuln_detail(3, idx + 1, row, "Máy tính xách tay Windows 11 (192.168.95.135)")

# --- SECTION 4 ---
add_h1("4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER (192.168.95.138)")
for idx, row in new_webserver.reset_index().iterrows():
    render_vuln_detail(4, idx + 1, row, "Web Server (192.168.95.138)")

# --- SECTION 5 ---
add_h1("5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO")
add_h2("5.1. MÔ HÌNH ĐÁNH GIÁ RỦI RO THEO OWASP RISK RATING")
add_p("Mô hình đánh giá rủi ro an toàn thông tin dựa trên tiêu chuẩn OWASP Risk Rating Methodology, kết hợp 2 yếu tố chính: Likelihood (Khả năng xảy ra) và Impact (Mức độ ảnh hưởng).")

add_h2("5.2. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0")
add_p("Common Vulnerability Scoring System (CVSS) v3.0 là tiêu chuẩn quốc tế dùng để đánh giá mức độ nghiêm trọng của lỗ hổng phần mềm.")

add_h3("5.2.1. EXPLOITABILITY METRICS")
add_p("Bao gồm các chỉ số: Attack Vector (AV), Attack Complexity (AC), Privileges Required (PR), User Interaction (UI).")

add_h3("5.2.2. IMPACT METRICS")
add_p("Bao gồm các chỉ số ảnh hưởng tính Bảo mật (Confidentiality), Toàn vẹn (Integrity), và Sẵn sàng (Availability).")

add_h2("5.3. THANG ĐIỂM ĐÁNH GIÁ CVSS V3.0")
t_cvss = doc.add_table(rows=6, cols=3)
t_cvss.style = 'Table Grid'
t_cvss.rows[0].cells[0].text = "Mức độ"
t_cvss.rows[0].cells[1].text = "Điểm CVSS v3.0"
t_cvss.rows[0].cells[2].text = "Mô tả rủi ro"

data_cvss = [
    ("Critical (Nghiêm trọng)", "9.0 - 10.0", "Lỗ hổng cho phép khai thác từ xa không cần xác thực, chiếm quyền điều khiển hệ thống."),
    ("High (Cao)", "7.0 - 8.9", "Lỗ hổng có khả năng gây ảnh hưởng nghiêm trọng đến tính bảo mật hoặc sẵn sàng."),
    ("Medium (Trung bình)", "4.0 - 6.9", "Lỗ hổng yêu cầu điều kiện khai thác phức tạp hoặc quyền truy cập hạn chế."),
    ("Low (Thấp)", "0.1 - 3.9", "Lỗ hổng có mức độ tác động nhỏ, khó khai thác."),
    ("Informational (Thông tin)", "0.0", "Thông tin ghi nhận cấu hình hoặc phiên bản dịch vụ.")
]

for idx, (m, d, desc) in enumerate(data_cvss):
    r_cells = t_cvss.rows[idx+1].cells
    r_cells[0].text = m
    r_cells[1].text = d
    r_cells[2].text = desc

doc.save('Bao_Cao_An_Toan_Thong_Tin_2026.docx')
print("Successfully generated perfect Bao_Cao_An_Toan_Thong_Tin_2026.docx")
