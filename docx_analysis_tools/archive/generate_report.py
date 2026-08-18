#!/usr/bin/env python3
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import json
import os

try:
    old_df = pd.read_csv('runs/run_20260810_140029/output/vuln_validation_queue.csv')
except:
    old_df = pd.read_csv('runs/run_20260810_140029/output/vuln_raw.csv')
    
try:
    new_df = pd.read_csv('runs/run_20260810_192535/output/vuln_validation_queue.csv')
except:
    new_df = pd.read_csv('runs/run_20260810_192535/output/vuln_raw.csv')

# Filter for the specific targets
old_laptop = old_df[old_df['location'].astype(str).str.contains('192.168.95.135', na=False) | old_df['asset'].astype(str).str.contains('192.168.95.135', na=False)]
new_webserver = new_df[new_df['location'].astype(str).str.contains('192.168.95.138', na=False) | new_df['asset'].astype(str).str.contains('192.168.95.138', na=False)]

print(f"Old laptop findings: {len(old_laptop)}")
print(f"New webserver findings: {len(new_webserver)}")

doc = docx.Document('MAU REPORT.docx')

# Clear existing body but keep headers/footers/styles
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
for t in doc.tables:
    t._element.getparent().remove(t._element)
    
def add_heading(text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.bold = True

add_heading("DỊCH VỤ AN TOÀN THÔNG TIN", 1)
doc.add_paragraph("------------------------------------------", style='Normal')
p = doc.add_paragraph("BÁO CÁO NỘI DUNG\nĐÁNH GIÁ LỖ HỔNG BẢO MẬT", style='Heading 1')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading("1. THÔNG TIN DỰ ÁN", 1)
add_heading("1.1. MỤC ĐÍCH ĐÁNH GIÁ", 2)
doc.add_paragraph("Kiểm tra và đánh giá toàn diện hiện trạng an toàn thông tin đối với các trang thiết bị đầu cuối (máy tính xách tay) và máy chủ dịch vụ (Web Server) đang vận hành trong hệ thống mạng định kỳ 3 tháng/lần.", style='Normal')

add_heading("1.2. PHẠM VI THỰC HIỆN", 2)
doc.add_paragraph("- Đánh giá nội bộ (Internal Scan): Máy tính xách tay Windows 11 (IP: 192.168.95.135)", style='List Paragraph')
doc.add_paragraph("- Đánh giá bên ngoài (External Scan - Internet Simulation): Hệ thống Web Server (IP: 192.168.95.138)", style='List Paragraph')

add_heading("2. BÁO CÁO TỔNG QUÁT", 1)
doc.add_paragraph("Quá trình đánh giá được phân loại rõ ràng theo chuẩn CVE/CVSS (Critical, High, Medium, Low).", style='Normal')

add_heading("2.1. TÌNH HÌNH LỖ HỔNG (CHU KỲ HIỆN TẠI)", 2)
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
headers = ['Đối tượng (Target)', 'Critical', 'High', 'Medium', 'Low', 'Info']
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].font.bold = True

# Calculate new webserver stats
nw_crit = len(new_webserver[new_webserver['severity'].str.upper() == 'CRITICAL'])
nw_high = len(new_webserver[new_webserver['severity'].str.upper() == 'HIGH'])
nw_med = len(new_webserver[new_webserver['severity'].str.upper() == 'MEDIUM'])
nw_low = len(new_webserver[new_webserver['severity'].str.upper() == 'LOW'])
nw_info = len(new_webserver[new_webserver['severity'].str.upper().isin(['INFORMATIONAL', 'INFO'])])

row_cells = table.add_row().cells
row_cells[0].text = 'Web Server (192.168.95.138)'
row_cells[1].text = str(nw_crit)
row_cells[2].text = str(nw_high)
row_cells[3].text = str(nw_med)
row_cells[4].text = str(nw_low)
row_cells[5].text = str(nw_info)

add_heading("2.2. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA/TRENDING)", 2)
doc.add_paragraph("So sánh bề mặt tấn công giữa 2 chu kỳ rà quét cho thấy sự dịch chuyển đáng kể của rủi ro:", style='Normal')
doc.add_paragraph("- Chu kỳ 3 tháng trước: Rủi ro tập trung vào thiết bị đầu cuối Máy tính xách tay Windows 11 (IP: 192.168.95.135) với các lỗ hổng cấp mạng (NetBIOS, SMB) ở mức High/Medium.", style='List Paragraph')
doc.add_paragraph(f"- Chu kỳ hiện tại: Phát hiện số lượng lớn lỗ hổng tràn ngập trên Web Server (IP: 192.168.95.138) với {nw_crit} lỗi Critical và {nw_high} lỗi High. Sự gia tăng bất thường này cho thấy Web Server đang là điểm yếu nghiêm trọng nhất và cần ưu tiên vá lỗi ngay lập tức.", style='List Paragraph')

add_heading("3. BÁO CÁO CHI TIẾT VÀ HƯỚNG DẪN KHẮC PHỤC", 1)

def add_finding(idx, row, is_laptop=False):
    finding = str(row.get('finding_name', 'Unknown'))
    sev = str(row.get('severity', 'Unknown'))
    cve = str(row.get('cve', 'N/A'))
    if pd.isna(cve) or cve == 'nan': cve = 'N/A'
    cvss = str(row.get('cvss', 'N/A'))
    if pd.isna(cvss) or cvss == 'nan': cvss = 'N/A'
    sol = str(row.get('solution', 'Không có hướng dẫn.'))
    if pd.isna(sol) or sol == 'nan' or sol == '': sol = str(row.get('scanner_solution', 'Không có hướng dẫn.'))
    
    add_heading(f"3.{idx}. {finding.upper()}", 2)
    
    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    t.cell(0,0).text = 'Mức độ (Severity)'
    p = t.cell(0,1).paragraphs[0]
    r = p.add_run(sev.upper())
    r.font.bold = True
    if sev.upper() in ['CRITICAL', 'HIGH']:
        r.font.color.rgb = RGBColor(255, 0, 0)
        
    t.cell(1,0).text = 'CVE / CVSS'
    t.cell(1,1).text = f"{cve} / {cvss}"
    
    t.cell(2,0).text = 'Đối tượng ảnh hưởng'
    t.cell(2,1).text = 'Máy tính xách tay Windows 11 (192.168.95.135)' if is_laptop else 'Web Server (192.168.95.138)'
    
    t.cell(3,0).text = 'Hướng dẫn khắc phục (Remediation)'
    t.cell(3,1).text = sol

counter = 1
add_heading("CHI TIẾT LỖ HỔNG - WEB SERVER (192.168.95.138)", 3)
for _, row in new_webserver.iterrows():
    add_finding(counter, row, False)
    counter += 1

add_heading("CHI TIẾT LỖ HỔNG - MÁY TÍNH XÁCH TAY (192.168.95.135) [GHI NHẬN TỪ 3 THÁNG TRƯỚC]", 3)
for _, row in old_laptop.iterrows():
    add_finding(counter, row, True)
    counter += 1

add_heading("4. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO", 1)
add_heading("4.1. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0", 2)
doc.add_paragraph("Common Vulnerability Scoring System (CVSS) là một tiêu chuẩn công nghiệp mở và miễn phí dùng để đánh giá mức độ nghiêm trọng của một lỗ hổng an ninh hệ thống máy tính...", style='Normal')

doc.save('Bao_Cao_An_Toan_Thong_Tin_2026.docx')
print("Successfully generated Bao_Cao_An_Toan_Thong_Tin_2026.docx")
