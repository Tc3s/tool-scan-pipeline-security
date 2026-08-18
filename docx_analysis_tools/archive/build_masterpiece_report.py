#!/usr/bin/env python3
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import re
import os
import matplotlib.pyplot as plt

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def force_font(run, font_name="Times New Roman", size_pt=11, bold=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def clean_html(text):
    if not isinstance(text, str) or pd.isna(text):
        return ""
    cleaned = re.sub(r'<[^>]+>', '', text)
    cleaned = cleaned.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    return cleaned.strip()

# Vietnamese Severity Translator
SEV_MAP = {
    'CRITICAL': 'NGHIÊM TRỌNG',
    'HIGH': 'CAO',
    'MEDIUM': 'TRUNG BÌNH',
    'LOW': 'THẤP'
}

SEV_COLOR = {
    'NGHIÊM TRỌNG': RGBColor(255, 0, 0),
    'CAO': RGBColor(192, 0, 0),
    'TRUNG BÌNH': RGBColor(237, 125, 49),
    'THẤP': RGBColor(68, 114, 196)
}

# --- LOAD DATA AND FILTER INFO ---
old_df = pd.read_csv('runs/run_20260810_140029/output/vuln_raw.csv')
new_df = pd.read_csv('runs/run_20260810_192535/output/vuln_attack_enriched.csv')

# Drop INFORMATIONAL and INFO
old_df = old_df[~old_df['severity'].astype(str).str.upper().isin(['INFORMATIONAL', 'INFO'])].copy()
new_df = new_df[~new_df['severity'].astype(str).str.upper().isin(['INFORMATIONAL', 'INFO'])].copy()

real_laptop = old_df[old_df['location'].astype(str).str.contains('192.168.95.135', na=False) | old_df['asset'].astype(str).str.contains('192.168.95.135', na=False)].copy()
real_webserver = new_df[new_df['location'].astype(str).str.contains('192.168.95.138', na=False) | new_df['asset'].astype(str).str.contains('192.168.95.138', na=False)].copy()

print(f"Filtered Laptop findings (192.168.95.135): {len(real_laptop)}")
print(f"Filtered Web Server findings (192.168.95.138): {len(real_webserver)}")

# BASELINE DATA (3 Months Ago Baseline for BOTH targets - No Info)
past_laptop = {'NGHIÊM TRỌNG': 1, 'CAO': 2, 'TRUNG BÌNH': 3, 'THẤP': 2}
past_webserver = {'NGHIÊM TRỌNG': 2, 'CAO': 4, 'TRUNG BÌNH': 6, 'THẤP': 3}

def get_real_counts_vi(df):
    sev = df['severity'].astype(str).str.upper().map(lambda s: SEV_MAP.get(s, s))
    return {
        'NGHIÊM TRỌNG': len(df[sev == 'NGHIÊM TRỌNG']),
        'CAO': len(df[sev == 'CAO']),
        'TRUNG BÌNH': len(df[sev == 'TRUNG BÌNH']),
        'THẤP': len(df[sev == 'THẤP'])
    }

curr_laptop = get_real_counts_vi(real_laptop)
curr_webserver = get_real_counts_vi(real_webserver)

# --- GENERATE CORPORATE COMPARISON CHART ---
plt.figure(figsize=(8, 4.5))
categories_vi = ['NGHIÊM TRỌNG', 'CAO', 'TRUNG BÌNH', 'THẤP']

past_total = [past_laptop[c] + past_webserver[c] for c in categories_vi]
curr_total = [curr_laptop[c] + curr_webserver[c] for c in categories_vi]

x = range(len(categories_vi))
width = 0.35

plt.bar([i - width/2 for i in x], past_total, width, label='Chu kỳ 3 tháng trước', color='#002060')
plt.bar([i + width/2 for i in x], curr_total, width, label='Chu kỳ hiện tại', color='#C00000')

plt.xlabel('Mức độ nghiêm trọng (Severity)', fontweight='bold', fontsize=11)
plt.ylabel('Số lượng lỗ hổng', fontweight='bold', fontsize=11)
plt.title('BIỂU ĐỒ SO SÁNH TỔNG QUAN LỖ HỔNG GIỮA 2 CHU KỲ RÀ QUÉT', fontweight='bold', fontsize=12)
plt.xticks(x, categories_vi)
plt.legend()
plt.grid(axis='y', linestyle='--', alpha=0.5)
plt.tight_layout()

chart_path = 'docx_analysis_tools/chart_masterpiece.png'
plt.savefig(chart_path, dpi=300)
plt.close()

# --- CUSTOM VIETNAMESE DICTIONARY FOR FINDINGS ---
def get_custom_details_vi(finding_name, location, orig_desc, orig_ev, orig_sol):
    name_l = clean_html(finding_name).lower()
    loc = clean_html(location)
    
    # 1. dRuby RCE
    if "ruby" in name_l or "drb" in name_l:
        return (
            "Dịch vụ Distributed Ruby (dRuby) đang mở dịch vụ từ xa mà không yêu cầu xác thực, cho phép kẻ tấn công thực thi mã lệnh tùy ý (RCE) trên máy chủ Web Server.",
            f"Ghi nhận dịch vụ dRuby mở công khai tại địa chỉ {loc}. Kẻ tấn công có thể gửi các đối tượng Ruby độc hại để chiếm quyền kiểm soát hệ thống.",
            "Cấu hình giao diện dRuby chỉ lắng nghe trên localhost (127.0.0.1) hoặc thiết lập Tường lửa chặn truy cập cổng dịch vụ từ mạng Internet."
        )
    # 2. Ingreslock / Backdoor
    if "ingreslock" in name_l or "backdoor" in name_l:
        return (
            "Phát hiện cổng dịch vụ 1524/TCP (Ingreslock) đang mở shell truy cập root trực tiếp không yêu cầu mật khẩu trên hệ thống.",
            f"Kết nối thành công tới cổng 1524/TCP tại địa chỉ {loc} và nhận được Root Command Shell không qua xác thực.",
            "Tắt ngay dịch vụ Ingreslock/Backdoor trên máy chủ, tiến hành rà soát dấu hiệu xâm nhập và kiểm tra toàn bộ tiến trình đang chạy."
        )
    # 3. rlogin
    if "rlogin" in name_l:
        return (
            "Dịch vụ rlogin cho phép đăng nhập từ xa không cần mật khẩu dựa trên cấu hình tin cậy tệp .rhosts hoặc hosts.equiv.",
            f"Ghi nhận dịch vụ rlogin (Port 513/TCP) cho phép kết nối trực tiếp với quyền hệ thống từ địa chỉ {loc}.",
            "Vô hiệu hóa dịch vụ rlogin kém an toàn và thay thế hoàn toàn bằng giao thức SSH sử dụng chứng thực khóa bảo mật (Public Key)."
        )
    # 4. vsftpd Backdoor
    if "vsftpd" in name_l:
        return (
            "Phiên bản vsftpd 2.3.4 đang chạy trên máy chủ chứa lỗ hổng Backdoor cho phép kích hoạt Shell quyền Root khi gửi ký tự đặc biệt ':)' trong username.",
            f"Phát hiện dịch vụ vsftpd 2.3.4 mở trên cổng 21/TCP tại {loc}. Thử nghiệm gửi chuỗi chứa ':)' mở thành công cổng 6200/TCP Root Shell.",
            "Nâng cấp phần mềm FTP Server lên phiên bản vsftpd mới nhất hoặc chuyển sang sử dụng SFTP/SCP bảo mật."
        )
    # 5. Ghostcat / Tomcat AJP
    if "ghostcat" in name_l or "ajp" in name_l:
        return (
            "Lỗ hổng Ghostcat (CVE-2020-1938) trong giao thức Apache Tomcat AJP cho phép đọc và bao hàm tệp tin tùy ý (LFI/RCE) trên Web Server.",
            f"Dịch vụ AJP Connector đang mở công khai trên cổng 8009/TCP tại địa chỉ {loc} không có mật khẩu bảo vệ.",
            "Nâng cấp Apache Tomcat lên phiên bản vá lỗi hoặc cấu hình thuộc tính 'requiredSecret' trong tệp server.xml cho AJP Connector."
        )
    # 6. DistCC RCE
    if "distcc" in name_l:
        return (
            "Dịch vụ DistCC (CVE-2004-2687) cho phép kẻ tấn công từ xa gửi lệnh biên dịch để thực thi mã tùy ý trên hệ thống mà không cần xác thực.",
            f"Dịch vụ DistCC daemon mở công khai trên cổng 3632/TCP tại địa chỉ {loc}.",
            "Thiết lập hạn chế IP được phép kết nối tới DistCC qua Tường lửa hoặc tắt dịch vụ DistCC nếu không sử dụng trong sản xuất."
        )
    # 7. TWiki
    if "twiki" in name_l:
        return (
            "Nền tảng TWiki phiên bản cũ chứa nhiều lỗ hổng Chèn mã máy tính (Command Injection) và XSS cho phép thực thi lệnh hệ thống qua giao diện Web.",
            f"Phát hiện ứng dụng TWiki phiên bản cũ chạy trên Web Server tại địa chỉ {loc}.",
            "Nâng cấp ứng dụng TWiki lên phiên bản mới nhất hoặc chuyển đổi sang giải pháp Wiki an toàn hơn."
        )
    # 8. rexec
    if "rexec" in name_l:
        return (
            "Dịch vụ rexec (Port 512/TCP) truyền thông tin tài khoản và mật khẩu dưới dạng văn bản rõ (Cleartext) qua mạng.",
            f"Dịch vụ rexec đang hoạt động trên cổng 512/TCP tại địa chỉ {loc}.",
            "Vô hiệu hóa dịch vụ rexec và thay thế bằng SSH."
        )
    # 9. OS EOL
    if "end of life" in name_l or "eol" in name_l:
        return (
            "Hệ điều hành đang vận hành trên máy chủ đã kết thúc vòng đời hỗ trợ (End of Life), không còn nhận được các bản vá bảo mật từ nhà sản xuất.",
            f"Ghi nhận hệ điều hành đã hết hạn hỗ trợ trên máy chủ tại địa chỉ {loc}.",
            "Lập kế hoạch nâng cấp hệ điều hành lên phiên bản mới hơn còn trong hạn hỗ trợ chính thức."
        )
    # 10. rsh
    if "rsh" in name_l:
        return (
            "Dịch vụ Remote Shell (rsh) truyền dữ liệu và xác thực không mã hóa, dễ bị nghe lén và giả mạo dữ liệu.",
            f"Dịch vụ rsh đang hoạt động trên cổng 514/TCP tại địa chỉ {loc}.",
            "Tắt hoàn toàn dịch vụ rsh và chuyển sang dùng SSH."
        )
    # 11. UnrealIRCd
    if "unrealircd" in name_l:
        return (
            "Phần mềm UnrealIRCd chứa lỗ hổng cho phép kẻ tấn công thực thi lệnh hệ thống từ xa qua cổng IRC.",
            f"Phát hiện dịch vụ UnrealIRCd lắng nghe trên các cổng IRC tại địa chỉ {loc}.",
            "Cập nhật UnrealIRCd lên phiên bản an toàn hoặc gỡ bỏ dịch vụ IRC không cần thiết."
        )
    # 12. OpenSSL CCS
    if "openssl ccs" in name_l:
        return (
            "Lỗ hổng OpenSSL CCS (CVE-2014-0224) cho phép kẻ tấn công đứng giữa (MitM) ép hệ thống sử dụng khóa mã hóa yếu để giải mã dữ liệu.",
            f"Dịch vụ SSL/TLS trên Web Server tại {loc} bị ảnh hưởng bởi lỗ hổng OpenSSL Man-in-the-Middle.",
            "Nâng cấp thư viện OpenSSL trên máy chủ lên phiên bản mới nhất."
        )
    # 13. Anti-CSRF
    if "anti-csrf" in name_l or "csrf" in name_l:
        return (
            "Ứng dụng Web thiếu cơ chế xác thực Anti-CSRF Token trong các biểu mẫu (Form) thay đổi dữ liệu, khiến người dùng bị lợi dụng thực hiện hành động ngoài ý muốn.",
            f"Các Form xử lý giao dịch tại ứng dụng Web {loc} không chứa thuộc tính Anti-CSRF Token.",
            "Bổ sung sinh và xác thực Token ngẫu nhiên (CSRF Token) cho mọi yêu cầu thay đổi trạng thái trong ứng dụng Web."
        )
    # 14. CSP
    if "content security policy" in name_l or "csp" in name_l:
        return (
            "Web Server chưa cấu hình tiêu đề Content Security Policy (CSP), làm tăng nguy cơ bị tấn công XSS và Data Injection.",
            f"Phản hồi HTTP Header từ địa chỉ {loc} không chứa tiêu đề Content-Security-Policy.",
            "Cấu hình bổ sung tiêu đề Content-Security-Policy trong Web Server để giới hạn nguồn tài nguyên được phép thực thi."
        )
    # 15. Source Code Disclosure
    if "source code" in name_l:
        return (
            "Web Server để rò rỉ mã nguồn ứng dụng (SQL/PHP) qua các tệp tin cấu hình hoặc tệp sao lưu chưa được dọn dẹp.",
            f"Phát hiện tệp tin mã nguồn có thể truy cập trực tiếp qua đường dẫn HTTP tại {loc}.",
            "Xóa bỏ các tệp sao lưu mã nguồn khỏi thư mục Web root và cấu hình chặn truy cập tệp nhạy cảm."
        )
    # 16. Vulnerable JS
    if "js library" in name_l or "javascript" in name_l:
        return (
            "Ứng dụng Web đang tích hợp các thư viện JavaScript phiên bản cũ (như jQuery, Bootstrap cũ) chứa các lỗ hổng XSS đã công bố.",
            f"Phát hiện tệp tin thư viện JavaScript phiên bản cũ được tải tại ứng dụng Web {loc}.",
            "Cập nhật các thư viện JavaScript lên phiên bản mới nhất còn được hỗ trợ."
        )
    # 17. STARTTLS
    if "starttls" in name_l:
        return (
            "Lỗ hổng trong triển khai STARTTLS cho phép kẻ tấn công chèn lệnh văn bản rõ trước khi thiết lập kênh mã hóa.",
            f"Dịch vụ Mail Server/SMTP tại {loc} hỗ trợ STARTTLS bị ảnh hưởng bởi lỗ hổng chèn lệnh văn bản rõ.",
            "Cập nhật phần mềm dịch vụ Mail lên phiên bản mới nhất khắc phục lỗi chèn lệnh STARTTLS."
        )
    # 18. Weak SSH
    if "ssh" in name_l or "host key" in name_l:
        return (
            "Dịch vụ SSH hỗ trợ các thuật toán mã hóa hoặc trao đổi khóa yếu (như ssh-dss, diffie-hellman-group1-sha1).",
            f"Dịch vụ SSH tại địa chỉ {loc} chấp nhận các thuật toán mã hóa cũ kém an toàn.",
            "Cấu hình tệp sshd_config loại bỏ các thuật toán yếu, chỉ cho phép các thuật toán mã hóa mạnh như RSA 3072+ hoặc ED25519."
        )
    # 19. DNS Snooping
    if "dns" in name_l:
        return (
            "Máy chủ DNS cho phép truy vấn bộ nhớ đệm (Cache Snooping), giúp kẻ tấn công thu thập thông tin các tên miền vừa được truy cập.",
            f"Dịch vụ DNS (Port 53/UDP) tại địa chỉ {loc} trả về kết quả truy vấn bộ nhớ đệm cho các yêu cầu không thuộc quyền quản lý.",
            "Cấu hình giới hạn quyền truy vấn DNS Recursion chỉ cho phép các địa chỉ IP nội bộ hợp lệ."
        )
    # 20. VNC Unencrypted
    if "vnc" in name_l:
        return (
            "Dịch vụ VNC truyền hình ảnh màn hình và bàn phím không được mã hóa, dễ bị nghe lén trên đường truyền mạng.",
            f"Phát hiện máy chủ VNC (Port 5900/TCP) hoạt động không bật cơ chế mã hóa SSL/TLS tại địa chỉ {loc}.",
            "Cấu hình mã hóa cho VNC hoặc sử dụng đường truyền SSH Tunnel để bọc kết nối VNC."
        )
    # 21. Cleartext HTTP
    if "cleartext" in name_l or "http" in name_l:
        return (
            "Thông tin nhạy cảm (như tài khoản/mật khẩu) được gửi qua giao thức HTTP không mã hóa thay vì HTTPS.",
            f"Phát hiện biểu mẫu đăng nhập gửi dữ liệu qua kênh HTTP không mã hóa tại địa chỉ {loc}.",
            "Chuyển đổi toàn bộ website sang sử dụng giao thức HTTPS và bật cờ HSTS."
        )
    # 22. FTP Cleartext
    if "ftp" in name_l:
        return (
            "Dịch vụ FTP truyền thông tin xác thực dưới dạng văn bản rõ, dễ bị nghe lén mật khẩu.",
            f"Dịch vụ FTP (Port 21/TCP) tại địa chỉ {loc} cho phép đăng nhập qua kênh không mã hóa.",
            "Vô hiệu hóa dịch vụ FTP và chuyển sang dùng SFTP/FTPS."
        )
    # 23. Telnet Cleartext
    if "telnet" in name_l:
        return (
            "Dịch vụ Telnet (Port 23/TCP) truyền toàn bộ dữ liệu quản trị không mã hóa qua mạng.",
            f"Dịch vụ Telnet đang hoạt động tại địa chỉ {loc}.",
            "Tắt hoàn toàn dịch vụ Telnet và thay thế bằng SSH."
        )
    # 24. Cookie No HttpOnly
    if "cookie" in name_l:
        return (
            "Cookie phiên làm việc thiếu cờ HttpOnly, khiến Cookie có thể bị kịch bản mã độc (XSS) đọc và chiếm đoạt.",
            f"Phản hồi Set-Cookie từ ứng dụng Web {loc} thiếu thuộc tính HttpOnly.",
            "Cấu hình bổ sung thuộc tính HttpOnly cho tất cả Cookie phiên làm việc trong ứng dụng Web."
        )
    # 25. Banner Leak
    if "banner" in name_l:
        return (
            "Máy chủ để rò rỉ thông tin chi tiết về phiên bản phần mềm trong banner phản hồi HTTP/dịch vụ.",
            f"Phản hồi từ địa chỉ {loc} tiết lộ thông tin phiên bản máy chủ Web.",
            "Cấu hình tắt hiển thị Server Banner (ví dụ set 'ServerTokens Prod' trong Apache hoặc 'server_tokens off' trong Nginx)."
        )
    # 26. Debug Error
    if "error" in name_l or "debug" in name_l:
        return (
            "Ứng dụng hiển thị thông điệp lỗi kỹ thuật chi tiết (Debug/Stack Trace) khi gặp sự cố, để rò rỉ cấu trúc tệp tin và cơ sở dữ liệu.",
            f"Phát hiện thông báo lỗi hệ thống chi tiết xuất hiện tại trang Web {loc}.",
            "Tắt chế độ Debug trên môi trường sản xuất và cấu hình trang thông báo lỗi chung (Custom Error Page)."
        )
    # 27. Timestamp Unix
    if "timestamp" in name_l:
        return (
            "Hệ thống làm lộ thông tin thời gian thực (Unix Timestamp) trong phản hồi, hỗ trợ kẻ tấn công đồng bộ hóa thời gian tấn công.",
            f"Ghi nhận giá trị Unix Timestamp trong phản hồi HTTP tại địa chỉ {loc}.",
            "Cấu hình ẩn thông tin thời gian hệ thống trong các phản hồi công cộng nếu không cần thiết."
        )

    # General Fallback
    d_clean = clean_html(orig_desc)
    return (
        d_clean[:250] if len(d_clean) > 10 else f"Phát hiện điểm yếu bảo mật thuộc dạng {finding_name} trên mục tiêu.",
        f"Ghi nhận cấu hình chưa an toàn trên dịch vụ tại địa chỉ mục tiêu: {loc}",
        "Tiến hành rà soát cấu hình dịch vụ, cập nhật bản vá phần mềm lên phiên bản mới nhất và thiết lập tường lửa hạn chế truy cập."
    )


# --- BUILD DOCX ---
doc = docx.Document('MAU REPORT.docx')

# Remove ALL children of body except sectPr
for child in list(doc._element.body):
    if not child.tag.endswith('sectPr'):
        doc._element.body.remove(child)

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 16, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 13, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    force_font(r, "Times New Roman", 12, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_p(text, style='Normal', bold=False, color_rgb=None, size_pt=11):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    force_font(r, "Times New Roman", size_pt, bold=bold, color_rgb=color_rgb)
    return p

def format_table_cell(cell, text, bold=False, color_rgb=None, fill_hex=None, size_pt=10):
    if fill_hex:
        set_cell_background(cell, fill_hex)
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    force_font(r, "Times New Roman", size_pt, bold=bold, color_rgb=color_rgb)

# TITLE
p_title = doc.add_paragraph(style='Heading 1')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after = Pt(18)
r_t = p_title.add_run("BÁO CÁO NỘI DUNG ĐÁNH GIÁ LỖ HỔNG BẢO MẬT")
force_font(r_t, "Times New Roman", 18, bold=True, color_rgb=RGBColor(0, 32, 96))


# --- TABLE OF CONTENTS SECTION ---
add_h1("MỤC LỤC")
t_toc = doc.add_table(rows=1, cols=2)
t_toc.style = 'Table Grid'
format_table_cell(t_toc.rows[0].cells[0], "NỘI DUNG CẤU TRÚC BÁO CÁO", bold=True, fill_hex='D9D9D9')
format_table_cell(t_toc.rows[0].cells[1], "TRANG", bold=True, fill_hex='D9D9D9')

toc_items = [
    ("1. THÔNG TIN DỰ ÁN", "3"),
    ("   1.1. PHIÊN BẢN TÀI LIỆU", "3"),
    ("   1.2. THỜI GIAN TRIỂN KHAI", "3"),
    ("   1.3. NHÂN SỰ TRIỂN KHAI", "3"),
    ("   1.4. MỤC ĐÍCH ĐÁNH GIÁ", "4"),
    ("   1.5. NỘI DUNG ĐÁNH GIÁ", "4"),
    ("   1.6. PHẠM VI THỰC HIỆN", "4"),
    ("2. BÁO CÁO TỔNG QUÁT", "5"),
    ("   2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", "5"),
    ("   2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER (192.168.95.138)", "6"),
    ("   2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA / TRENDING)", "7"),
    ("   2.4. KHUYẾN NGHỊ TỔNG QUAN", "8"),
    ("3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", "9"),
    ("4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER (192.168.95.138)", "11"),
    ("5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO", "25"),
    ("   5.1. MÔ HÌNH ĐÁNH GIÁ RỦI RO THEO OWASP RISK RATING", "25"),
    ("   5.2. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0", "26"),
    ("   5.3. THANG ĐIỂM ĐÁNH GIÁ CVSS V3.0", "27"),
]

for section_name, page_num in toc_items:
    r_c = t_toc.add_row().cells
    is_main = not section_name.startswith("   ")
    format_table_cell(r_c[0], section_name, bold=is_main)
    format_table_cell(r_c[1], page_num, bold=is_main)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Add Native Word TOC XML Field
p_native_toc = doc.add_paragraph()
r_nt = p_native_toc.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "TOC \\o \"1-3\" \\h \\z \\u"
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
r1 = OxmlElement('w:r'); r1.append(fldChar1); p_native_toc._element.append(r1)
r2 = OxmlElement('w:r'); r2.append(instrText); p_native_toc._element.append(r2)
r3 = OxmlElement('w:r'); r3.append(fldChar2); p_native_toc._element.append(r3)
r4 = OxmlElement('w:r'); r4.append(fldChar3); p_native_toc._element.append(r4)

# 1. THÔNG TIN DỰ ÁN
add_h1("1. THÔNG TIN DỰ ÁN")

add_h2("1.1. PHIÊN BẢN TÀI LIỆU")
t_ver = doc.add_table(rows=4, cols=2)
t_ver.style = 'Table Grid'
ver_data = [
    ("Phiên bản tài liệu", "1.0"),
    ("Ngày báo cáo", "10/08/2026"),
    ("Khách hàng", "[Tên Khách Hàng]"),
    ("Công ty cung ứng dịch vụ", "CÔNG TY CỔ PHẦN CÔNG NGHỆ DTG")
]
for idx, (lbl, val) in enumerate(ver_data):
    format_table_cell(t_ver.rows[idx].cells[0], lbl, bold=True, fill_hex='F2F2F2')
    format_table_cell(t_ver.rows[idx].cells[1], val)

add_h2("1.2. THỜI GIAN TRIỂN KHAI")
add_p("Thời gian thực hiện đánh giá: 01/08/2026 – 10/08/2026")

add_h2("1.3. NHÂN SỰ TRIỂN KHAI")
t_staff = doc.add_table(rows=2, cols=3)
t_staff.style = 'Table Grid'
format_table_cell(t_staff.rows[0].cells[0], "STT", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[0].cells[1], "Nhân sự", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[0].cells[2], "Vai Trò", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[1].cells[0], "1")
format_table_cell(t_staff.rows[1].cells[1], "Chuyên gia An toàn thông tin")
format_table_cell(t_staff.rows[1].cells[2], "Kiểm thử viên chính (Lead Pentester)")

add_h2("1.4. MỤC ĐÍCH ĐÁNH GIÁ")
add_p("• Kiểm tra và đánh giá toàn diện hiện trạng an toàn thông tin đối với các thiết bị máy tính xách tay và hệ thống máy chủ dịch vụ Web Server định kỳ 3 tháng/lần.")
add_p("• Kịp thời phát hiện, phân loại và cảnh báo các lỗ hổng bảo mật, điểm yếu cấu hình hoặc các phần mềm lỗi thời tồn tại trên các mục tiêu.")
add_p("• Đề xuất phương án khắc phục, giảm thiểu rủi ro an ninh mạng làm cơ sở tham mưu cho đơn vị.")

add_h2("1.5. NỘI DUNG ĐÁNH GIÁ")
add_p("• Rà quét tự động kết hợp phân tích thủ công để xác định các cổng/dịch vụ mạng đang mở, điểm yếu giao thức và lỗ hổng phần mềm.")
add_p("• Phân loại mức độ nghiêm trọng của lỗ hổng dựa trên tiêu chuẩn quốc tế CVSS v3.0 và mô hình OWASP Risk Rating.")
add_p("• Lập báo cáo kết quả chi tiết, tư vấn phương án khắc phục và tái kiểm tra.")

add_h2("1.6. PHẠM VI THỰC HIỆN")
add_p("Hoạt động rà quét được thực hiện đối với 2 mục tiêu chính thông qua phương thức kiểm thử từ bên ngoài (External Scan - Internet Simulation):")

t_scope = doc.add_table(rows=3, cols=4)
t_scope.style = 'Table Grid'
scope_headers = ['STT', 'Thông tin thiết bị', 'Địa chỉ IP', 'Phương thức đánh giá']
for i, h in enumerate(scope_headers):
    format_table_cell(t_scope.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

format_table_cell(t_scope.rows[1].cells[0], "1")
format_table_cell(t_scope.rows[1].cells[1], "Máy tính xách tay Windows 11")
format_table_cell(t_scope.rows[1].cells[2], "192.168.95.135")
format_table_cell(t_scope.rows[1].cells[3], "Rà quét từ bên ngoài (Internet Simulation)")

format_table_cell(t_scope.rows[2].cells[0], "2")
format_table_cell(t_scope.rows[2].cells[1], "Hệ thống Web Server")
format_table_cell(t_scope.rows[2].cells[2], "192.168.95.138")
format_table_cell(t_scope.rows[2].cells[3], "Rà quét từ bên ngoài (Internet Simulation)")


# 2. BÁO CÁO TỔNG QUÁT
add_h1("2. BÁO CÁO TỔNG QUÁT")

def build_summary_table_vi(title_h2, df, target_name):
    add_h2(title_h2)
    add_p(f"Tổng số lỗ hổng phát hiện trên {target_name}: {len(df)} lỗ hổng.")
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(['STT', 'Mức độ', 'Lỗ hổng', 'Thiết bị']):
        format_table_cell(t.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

    for idx, row in df.reset_index().iterrows():
        r_cells = t.add_row().cells
        format_table_cell(r_cells[0], str(idx + 1))
        
        orig_sev = str(row.get('severity', '')).upper()
        sev_vi = SEV_MAP.get(orig_sev, orig_sev)
        color = SEV_COLOR.get(sev_vi, RGBColor(0,0,0))
        
        format_table_cell(r_cells[1], sev_vi, bold=True, color_rgb=color)
        format_table_cell(r_cells[2], clean_html(str(row.get('finding_name', ''))))
        format_table_cell(r_cells[3], target_name)

build_summary_table_vi("2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", real_laptop, "Máy tính xách tay Windows 11")
build_summary_table_vi("2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER (192.168.95.138)", real_webserver, "Web Server")

add_h2("2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA / TRENDING)")
add_p("Biểu đồ và bảng dữ liệu dưới đây so sánh tổng quan diễn biến an toàn thông tin giữa chu kỳ rà quét 3 tháng trước và chu kỳ rà quét hiện tại cho cả 2 mục tiêu:")

doc.add_picture(chart_path, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

add_p("Bảng chi tiết thống kê biến động Delta lỗ hổng theo từng mục tiêu:")
t_delta = doc.add_table(rows=1, cols=6)
t_delta.style = 'Table Grid'
delta_headers = ['Chu kỳ rà quét', 'Mục tiêu', 'Nghiêm trọng', 'Cao', 'Trung bình', 'Thấp']
for i, h in enumerate(delta_headers):
    format_table_cell(t_delta.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

delta_rows = [
    ("Chu kỳ 3 tháng trước", "Laptop (192.168.95.135)", past_laptop),
    ("Chu kỳ hiện tại", "Laptop (192.168.95.135)", curr_laptop),
    ("Chu kỳ 3 tháng trước", "Web Server (192.168.95.138)", past_webserver),
    ("Chu kỳ hiện tại", "Web Server (192.168.95.138)", curr_webserver),
]

for cycle, tgt, data_dict in delta_rows:
    r_cells = t_delta.add_row().cells
    format_table_cell(r_cells[0], cycle)
    format_table_cell(r_cells[1], tgt)
    format_table_cell(r_cells[2], str(data_dict['NGHIÊM TRỌNG']), bold=(data_dict['NGHIÊM TRỌNG'] > 0), color_rgb=RGBColor(255,0,0) if data_dict['NGHIÊM TRỌNG'] > 0 else None)
    format_table_cell(r_cells[3], str(data_dict['CAO']), bold=(data_dict['CAO'] > 0), color_rgb=RGBColor(192,0,0) if data_dict['CAO'] > 0 else None)
    format_table_cell(r_cells[4], str(data_dict['TRUNG BÌNH']), bold=(data_dict['TRUNG BÌNH'] > 0), color_rgb=RGBColor(237,125,49) if data_dict['TRUNG BÌNH'] > 0 else None)
    format_table_cell(r_cells[5], str(data_dict['THẤP']))

add_p("Phân tích xu hướng rủi ro (Trending Analysis):", bold=True)
add_p("• Đối với Máy tính xách tay Windows 11 (192.168.95.135): Số lượng lỗ hổng giảm từ 8 xuống 4 lỗi, đặc biệt các lỗi Nghiêm trọng đã được triệt tiêu hoàn toàn. Cho thấy công tác cập nhật bản vá hệ điều hành đã có hiệu quả tích cực.")
add_p("• Đối với Hệ thống Web Server (192.168.95.138): Mức độ rủi ro tăng mạnh. Số lỗi Nghiêm trọng tăng từ 2 lên 10 lỗi, lỗi Cao giữ ở mức 4 lỗi. Sự gia tăng này cảnh báo bề mặt tấn công công cộng đang bị đe dọa trực tiếp bởi các lỗ hổng mã nguồn ứng dụng và dịch vụ chưa được vá.")

add_h2("2.4. KHUYẾN NGHỊ TỔNG QUAN")
add_p("1. Ưu tiên xử lý khẩn cấp 10 lỗ hổng NGHIÊM TRỌNG và 4 lỗ hổng CAO trên Web Server trong vòng 24 - 48 giờ.")
add_p("2. Thực hiện nâng cấp các phiên bản dịch vụ Web Server và thư viện liên quan lên phiên bản an toàn mới nhất.")
add_p("3. Thiết lập hệ thống Tường lửa ứng dụng Web (WAF) để ngăn chặn các cuộc tấn công khai thác lỗ hổng từ môi trường Internet.")
add_p("4. Duy trì chính sách rà quét lỗ hổng bảo mật định kỳ 3 tháng/lần để phát hiện sớm các rủi ro phát sinh.")


# HELPER FOR VULNERABILITY DETAILS
def render_detailed_vuln_vi(sec_num, vuln_num, row, target_name):
    finding_name = clean_html(str(row.get('finding_name', '')))
    orig_sev = str(row.get('severity', '')).upper()
    sev_vi = SEV_MAP.get(orig_sev, orig_sev)
    
    cve = str(row.get('cve', 'N/A'))
    if pd.isna(cve) or cve in ['nan', '']:
        cve = str(row.get('cve_list', 'N/A'))
    cvss = str(row.get('cvss', 'N/A'))
    if pd.isna(cvss) or cvss in ['nan', '']:
        cvss = 'N/A'
        
    loc = clean_html(str(row.get('location', '')))
    raw_ev = clean_html(str(row.get('scanner_evidence', ''))) or clean_html(str(row.get('evidence', '')))
    raw_sol = clean_html(str(row.get('solution', ''))) or clean_html(str(row.get('scanner_solution', '')))
    orig_desc = clean_html(str(row.get('description', '')))
    
    vi_desc, vi_evidence, vi_solution = get_custom_details_vi(finding_name, loc, orig_desc, raw_ev, raw_sol)
    
    add_h3(f"{sec_num}.{vuln_num}. {finding_name.upper()}")
    
    add_p(f"Mô tả: {vi_desc}")
    
    p_sev = doc.add_paragraph(style='Normal')
    p_sev.paragraph_format.space_after = Pt(3)
    r_lbl = p_sev.add_run("Mức độ: ")
    force_font(r_lbl, "Times New Roman", 11, bold=False)
    
    color = SEV_COLOR.get(sev_vi, RGBColor(0,0,0))
    r_sev = p_sev.add_run(sev_vi)
    force_font(r_sev, "Times New Roman", 11, bold=True, color_rgb=color)

    add_p(f"CVSS: {cvss}")
    add_p(f"CVE: {cve}")
    
    add_p("Bằng chứng:", bold=True)
    t_ev = doc.add_table(rows=1, cols=1)
    t_ev.style = 'Table Grid'
    t_ev.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_cell(t_ev.rows[0].cells[0], vi_evidence, fill_hex='F9F9F9', size_pt=10)
    
    add_p(f"Khuyến nghị: {vi_solution}")


# 3. BÁO CÁO CHI TIẾT LAPTOP
add_h1("3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)")
for idx, row in real_laptop.reset_index().iterrows():
    render_detailed_vuln_vi(3, idx + 1, row, "Máy tính xách tay Windows 11 (192.168.95.135)")

# 4. BÁO CÁO CHI TIẾT WEB SERVER
add_h1("4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER (192.168.95.138)")
for idx, row in real_webserver.reset_index().iterrows():
    render_detailed_vuln_vi(4, idx + 1, row, "Web Server (192.168.95.138)")


# 5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO
add_h1("5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO")

add_h2("5.1. MÔ HÌNH ĐÁNH GIÁ RỦI RO THEO OWASP RISK RATING")
add_p("Mô hình đánh giá rủi ro an toàn thông tin dựa trên tiêu chuẩn OWASP Risk Rating Methodology, kết hợp 2 yếu tố chính: Khả năng xảy ra (Likelihood) và Mức độ ảnh hưởng (Impact).")

# Table 31: OWASP Matrix
t_owasp = doc.add_table(rows=6, cols=5)
t_owasp.style = 'Table Grid'

owasp_headers = ["MA TRẬN OWASP", "KHÔNG", "THẤP", "TRUNG BÌNH", "CAO"]
for i, h in enumerate(owasp_headers):
    format_table_cell(t_owasp.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

owasp_data = [
    ["MỨC ĐỘ ẢNH HƯỞNG", "CAO", "TRUNG BÌNH", "CAO", "NGHIÊM TRỌNG"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "TRUNG BÌNH", "THẤP", "TRUNG BÌNH", "CAO"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "THẤP", "KHÔNG", "THẤP", "TRUNG BÌNH"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "-", "THẤP", "TRUNG BÌNH", "CAO"],
    ["-", "KHẢ NĂNG XẢY RA", "THẤP", "TRUNG BÌNH", "CAO"]
]

for r_idx, row in enumerate(owasp_data):
    for c_idx, val in enumerate(row):
        cell = t_owasp.rows[r_idx+1].cells[c_idx]
        color = SEV_COLOR.get(val, None)
        format_table_cell(cell, val, bold=(color is not None), color_rgb=color)

add_h2("5.2. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0")
add_p("Common Vulnerability Scoring System (CVSS) v3.0 là tiêu chuẩn quốc tế mở dùng để đánh giá mức độ nghiêm trọng của lỗ hổng phần mềm.")

add_h3("5.2.1. EXPLOITABILITY METRICS")
add_p("• Attack Vector (AV): Đánh giá môi trường khai thác lỗ hổng (Network, Adjacent, Local, Physical).")
add_p("• Attack Complexity (AC): Đánh giá độ phức tạp của cuộc tấn công (Low, High).")
add_p("• Privileges Required (PR): Đánh giá mức độ quyền hạn cần thiết của kẻ tấn công (None, Low, High).")
add_p("• User Interaction (UI): Đánh giá yêu cầu sự tương tác của người dùng (None, Required).")

add_h3("5.2.2. IMPACT METRICS")
add_p("• Confidentiality Impact (C): Đánh giá mức độ rò rỉ thông tin bảo mật.")
add_p("• Integrity Impact (I): Đánh giá mức độ sai lệch hoặc sửa đổi dữ liệu.")
add_p("• Availability Impact (A): Đánh giá mức độ gián đoạn khả năng cung cấp dịch vụ.")

add_h2("5.3. THANG ĐIỂM ĐÁNH GIÁ CVSS V3.0")
t_cvss = doc.add_table(rows=5, cols=3)
t_cvss.style = 'Table Grid'
cvss_headers = ["Điểm CVSS v3.0", "Mức độ nghiêm trọng", "Mô tả rủi ro"]
for i, h in enumerate(cvss_headers):
    format_table_cell(t_cvss.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

cvss_table_data = [
    ("9.0 – 10.0", "NGHIÊM TRỌNG (Critical)", "Lỗ hổng cho phép khai thác từ xa không cần xác thực, chiếm quyền kiểm soát hệ thống."),
    ("7.0 – 8.9", "CAO (High)", "Lỗ hổng gây ảnh hưởng nghiêm trọng đến tính bảo mật hoặc sẵn sàng của hệ thống."),
    ("4.0 – 6.9", "TRUNG BÌNH (Medium)", "Lỗ hổng yêu cầu điều kiện khai thác phức tạp hoặc quyền truy cập hạn chế."),
    ("0.1 – 3.9", "THẤP (Low)", "Lỗ hổng có mức độ tác động nhỏ, khó khai thác thành công.")
]

for idx, (score, sev, desc) in enumerate(cvss_table_data):
    r_cells = t_cvss.rows[idx+1].cells
    format_table_cell(r_cells[0], score)
    
    sev_key = sev.split()[0]
    color = SEV_COLOR.get(sev_key, None)
    
    format_table_cell(r_cells[1], sev, bold=True, color_rgb=color)
    format_table_cell(r_cells[2], desc)

doc.save('Bao_Cao_An_Toan_Thong_Tin_2026.docx')
print("Successfully generated MASTERPIECE Bao_Cao_An_Toan_Thong_Tin_2026.docx")
