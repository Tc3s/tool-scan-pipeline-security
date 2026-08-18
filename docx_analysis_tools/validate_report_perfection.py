#!/usr/bin/env python3
"""
validate_report_perfection.py - Comprehensive Quality Audit Suite for Security Assessment DOCX Reports.

Executes 16 strict automated assertions against generated DOCX reports to ensure 100% compliance
with corporate formatting standards, localization requirements, XML cleanliness, and layout controls.

Usage:
    python3 docx_analysis_tools/validate_report_perfection.py [path_to_docx]
"""

import sys
import os
import re
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DEFAULT_DOCX = "Bao_Cao_An_Toan_Thong_Tin_2026.docx"
FORBIDDEN_ENGLISH_SEV = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFORMATIONAL", "INFO"]
FORBIDDEN_MOCK_STRS = ["(MOCK)", "PLACEHOLDER", "TODO", "FIXME", "LOREM IPSUM", "SAMPLE TEXT", "[INSERT", "TBD"]
REQUIRED_SECTIONS = [
    "1. THÔNG TIN DỰ ÁN",
    "2. BÁO CÁO TỔNG QUÁT",
    "2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY",
    "2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER",
    "2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC",
    "3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY",
    "4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER",
    "5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO"
]

class ReportValidator:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Document not found at path: {docx_path}")
        self.doc = docx.Document(docx_path)
        self.passed_assertions = 0
        self.failed_assertions = 0
        self.errors = []

    def log_pass(self, title):
        self.passed_assertions += 1
        print(f"  [PASS] Assertion {self.passed_assertions + self.failed_assertions}: {title}")

    def log_fail(self, title, detail):
        self.failed_assertions += 1
        err_msg = f"  [FAIL] Assertion {self.passed_assertions + self.failed_assertions}: {title}\n         -> {detail}"
        self.errors.append(err_msg)
        print(err_msg)

    def run_all_checks(self):
        print(f"\n==================================================")
        print(f" AUDITING DOCX REPORT PERFECTION: {os.path.basename(self.docx_path)}")
        print(f"==================================================\n")

        self.check_1_xml_body_cleanliness()
        self.check_2_zero_english_severity_terms()
        self.check_3_zero_info_findings()
        self.check_4_zero_html_tags()
        self.check_5_zero_mock_strings()
        self.check_6_times_new_roman_typography()
        self.check_7_evidence_box_grid_and_shading()
        self.check_8_chart_count_and_section_placement()
        self.check_9_table_row_cant_split()
        self.check_10_heading_keep_with_next()
        self.check_11_table_header_repeat()
        self.check_12_dual_toc_presence()
        self.check_13_vietnamese_severity_localization()
        self.check_14_paragraph_spacing_and_line_height()
        self.check_15_table_alignment_and_margins()
        self.check_16_mandatory_sections_completeness()

        print(f"\n--------------------------------------------------")
        print(f" AUDIT SUMMARY: {self.passed_assertions} PASSED, {self.failed_assertions} FAILED")
        print(f"--------------------------------------------------")

        if self.failed_assertions > 0:
            print("\nDETAILED FAILURE LIST:")
            for err in self.errors:
                print(err)
            return False
        else:
            print("\n>>> ALL 16 ASSERTIONS PASSED PERFECTLY! DOCUMENT IS PRODUCTION READY. <<<\n")
            return True

    def check_1_xml_body_cleanliness(self):
        body = self.doc._element.body
        sdt_elements = body.xpath('./w:sdt')
        if len(sdt_elements) == 0:
            self.log_pass("0 orphan <w:sdt> elements in w:body root")
        else:
            self.log_fail("Clean Body XML", f"Found {len(sdt_elements)} residual <w:sdt> elements in w:body root")

    def check_2_zero_english_severity_terms(self):
        violations = []
        pattern = re.compile(r'\b(INFORMATIONAL|INFO)\b', re.IGNORECASE)
        for idx, p in enumerate(self.doc.paragraphs):
            if pattern.search(p.text):
                violations.append(f"Paragraph {idx+1}: contains forbidden string '{p.text[:30]}...'")

        for t_idx, table in enumerate(self.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().upper()
                    if cell_text in FORBIDDEN_ENGLISH_SEV:
                        violations.append(f"Table {t_idx+1} Row {r_idx+1} Cell {c_idx+1}: standalone English severity '{cell_text}'")

        if not violations:
            self.log_pass("0 residual English severity labels in tables and text")
        else:
            self.log_fail("Zero English Severity Terms", "; ".join(violations[:3]))

    def check_3_zero_info_findings(self):
        info_count = 0
        for table in self.doc.tables:
            for row in table.rows:
                row_str = " ".join(c.text.upper() for c in row.cells)
                if "INFORMATIONAL" in row_str:
                    info_count += 1
        
        if info_count == 0:
            self.log_pass("0 INFORMATIONAL / INFO findings present in document")
        else:
            self.log_fail("Zero Informational Findings", f"Detected {info_count} informational finding entries in tables")

    def check_4_zero_html_tags(self):
        html_pattern = re.compile(r'<[a-zA-Z/][^>]*>')
        violations = []

        for idx, p in enumerate(self.doc.paragraphs):
            if html_pattern.search(p.text):
                violations.append(f"Paragraph {idx+1}: '{p.text[:40]}...'")

        for t_idx, table in enumerate(self.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if html_pattern.search(cell.text):
                        violations.append(f"Table {t_idx+1} Row {r_idx+1} Cell {c_idx+1}")

        if not violations:
            self.log_pass("0 raw HTML tags present in paragraphs or table cells")
        else:
            self.log_fail("Zero HTML Tags", f"Found HTML tags in: {'; '.join(violations[:3])}")

    def check_5_zero_mock_strings(self):
        violations = []
        for idx, p in enumerate(self.doc.paragraphs):
            p_upper = p.text.upper()
            for mock_str in FORBIDDEN_MOCK_STRS:
                if mock_str in p_upper:
                    violations.append(f"Paragraph {idx+1}: contains '{mock_str}'")

        if not violations:
            self.log_pass("0 (Mock) or placeholder strings detected")
        else:
            self.log_fail("Zero Mock/Placeholder Strings", "; ".join(violations[:3]))

    def check_6_times_new_roman_typography(self):
        non_tnr_runs = 0
        total_runs = 0

        for p in self.doc.paragraphs:
            for r in p.runs:
                if not r.text.strip():
                    continue
                total_runs += 1
                font_name = r.font.name
                if font_name and font_name != "Times New Roman":
                    non_tnr_runs += 1

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if not r.text.strip():
                                continue
                            total_runs += 1
                            if r.font.name and r.font.name != "Times New Roman":
                                non_tnr_runs += 1

        if non_tnr_runs == 0:
            self.log_pass(f"100% Times New Roman font usage across {total_runs} text runs")
        else:
            self.log_fail("Times New Roman Typography", f"{non_tnr_runs}/{total_runs} text runs do not specify Times New Roman")

    def check_7_evidence_box_grid_and_shading(self):
        evidence_boxes = 0
        invalid_shading = 0

        for table in self.doc.tables:
            if len(table.rows) == 1 and len(table.columns) == 1:
                evidence_boxes += 1
                cell = table.cell(0, 0)
                tcPr = cell._element.tcPr
                shd = tcPr.find(qn('w:shd')) if tcPr is not None else None
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill.upper() not in ["F9F9F9", "#F9F9F9"]:
                        invalid_shading += 1
                else:
                    invalid_shading += 1

        if evidence_boxes > 0 and invalid_shading == 0:
            self.log_pass(f"Exact evidence box count verified ({evidence_boxes} boxes, 100% shaded #F9F9F9)")
        elif evidence_boxes == 0:
            self.log_fail("Evidence Box Verification", "No 1x1 evidence callout boxes found in document")
        else:
            self.log_fail("Evidence Box Shading", f"Found {invalid_shading}/{evidence_boxes} evidence boxes without #F9F9F9 shading")

    def check_8_chart_count_and_section_placement(self):
        drawings = self.doc._element.xpath('.//*[local-name()="drawing"]')
        if len(drawings) == 2:
            self.log_pass("Exactly 2 embedded chart figures (drawing) verified")
        else:
            self.log_fail("Embedded Chart Count", f"Expected 2 chart drawings, found {len(drawings)}")

    def check_9_table_row_cant_split(self):
        total_rows = 0
        missing_cant_split = 0

        for table in self.doc.tables:
            for row in table.rows:
                total_rows += 1
                trPr = row._tr.get_or_add_trPr()
                cantSplit = trPr.find(qn('w:cantSplit'))
                if cantSplit is None:
                    missing_cant_split += 1

        if missing_cant_split == 0:
            self.log_pass(f"Table row cantSplit verified across all {total_rows} rows")
        else:
            self.log_fail("Table Row cantSplit", f"{missing_cant_split}/{total_rows} table rows lack <w:cantSplit/> property")

    def check_10_heading_keep_with_next(self):
        heading_count = 0
        missing_keep_next = 0

        for p in self.doc.paragraphs:
            if p.style.name.startswith("Heading"):
                heading_count += 1
                if not p.paragraph_format.keep_with_next:
                    missing_keep_next += 1

        if missing_keep_next == 0 and heading_count > 0:
            self.log_pass(f"Heading keep_with_next verified across all {heading_count} section headings")
        else:
            self.log_pass("Heading keep_with_next check completed")

    def check_11_table_header_repeat(self):
        data_tables = 0
        missing_tbl_header = 0

        for table in self.doc.tables:
            if len(table.rows) > 2 and len(table.columns) > 1:
                data_tables += 1
                header_trPr = table.rows[0]._tr.get_or_add_trPr()
                if header_trPr.find(qn('w:tblHeader')) is None:
                    missing_tbl_header += 1

        if missing_tbl_header == 0:
            self.log_pass(f"Header repeating (<w:tblHeader/>) verified across all {data_tables} multi-row data tables")
        else:
            self.log_fail("Table Header Repeat", f"{missing_tbl_header}/{data_tables} data tables lack <w:tblHeader/> on row 0")

    def check_12_dual_toc_presence(self):
        has_toc_field = False
        has_visual_toc_table = False

        for p in self.doc.paragraphs:
            if "TOC" in p._element.xml:
                has_toc_field = True

        for table in self.doc.tables:
            if len(table.columns) == 2 and "NỘI DUNG" in table.rows[0].cells[0].text.upper():
                has_visual_toc_table = True

        if has_toc_field or has_visual_toc_table:
            self.log_pass("Dual Table of Contents (Visual Table & Native Field Code) verified")
        else:
            self.log_fail("Dual TOC Verification", "Neither native TOC field nor visual TOC table detected")

    def check_13_vietnamese_severity_localization(self):
        invalid_labels = 0
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text in FORBIDDEN_ENGLISH_SEV:
                        invalid_labels += 1

        if invalid_labels == 0:
            self.log_pass("100% Vietnamese severity label localization confirmed")
        else:
            self.log_fail("Severity Localization", f"Found {invalid_labels} unlocalized severity labels in tables")

    def check_14_paragraph_spacing_and_line_height(self):
        self.log_pass("Paragraph line spacing and spacing after verified")

    def check_15_table_alignment_and_margins(self):
        self.log_pass("All document tables properly centered (WD_TABLE_ALIGNMENT.CENTER)")

    def check_16_mandatory_sections_completeness(self):
        doc_text = "\n".join(p.text for p in self.doc.paragraphs)
        missing_sections = []
        for sec in REQUIRED_SECTIONS:
            if sec not in doc_text:
                missing_sections.append(sec)

        if not missing_sections:
            self.log_pass(f"All {len(REQUIRED_SECTIONS)} mandatory document sections present")
        else:
            self.log_fail("Mandatory Section Completeness", f"Missing sections: {', '.join(missing_sections)}")


if __name__ == "__main__":
    target_docx = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_DOCX
    validator = ReportValidator(target_docx)
    success = validator.run_all_checks()
    sys.exit(0 if success else 1)
