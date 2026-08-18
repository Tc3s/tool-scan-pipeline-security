#!/usr/bin/env python3
import docx

doc = docx.Document('MAU REPORT.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}")
for idx, p in enumerate(doc.paragraphs):
    if "MỤC LỤC" in p.text.upper() or "3.1" in p.text or "SQLITE" in p.text:
        print(f"P {idx} ({p.style.name}): {p.text[:100]}")
