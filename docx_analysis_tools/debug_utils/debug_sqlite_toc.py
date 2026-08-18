#!/usr/bin/env python3
import docx

doc = docx.Document('MAU REPORT.docx')

print("=== CHECKING SDT / TOC CONTROL ELEMENTS IN MAU REPORT.docx ===")
for elem in doc._element.body.xpath('.//*[local-name()="sdt" or local-name()="fldSimple" or local-name()="instrText"]'):
    print(f"Elem ({elem.tag}): {elem.text}")

# Also check generated document Bao_Cao_An_Toan_Thong_Tin_2026.docx for "SQLITE"
doc_gen = docx.Document('Bao_Cao_An_Toan_Thong_Tin_2026.docx')
print("\n=== CHECKING GENERATED DOCX FOR SQLITE ===")
sqlite_found = False
for p in doc_gen.paragraphs:
    if "SQLITE" in p.text.upper():
        print(f"Found SQLITE in paragraph: {p.text}")
        sqlite_found = True

for t_idx, t in enumerate(doc_gen.tables):
    for r in t.rows:
        for c in r.cells:
            if "SQLITE" in c.text.upper():
                print(f"Found SQLITE in Table {t_idx}: {c.text[:100]}")
                sqlite_found = True

if not sqlite_found:
    print("SQLITE not found in text/tables of generated docx! It must be in a Word Field / SDT block!")
