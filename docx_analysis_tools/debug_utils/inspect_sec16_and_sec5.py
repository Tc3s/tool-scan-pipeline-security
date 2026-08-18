#!/usr/bin/env python3
import docx

doc = docx.Document('MAU REPORT.docx')

print("=== SECTION 1.6 IN MAU REPORT.docx ===")
recording = False
for p in doc.paragraphs:
    if "1.6." in p.text:
        recording = True
    elif "2." in p.text and "BÁO CÁO TỔNG QUÁT" in p.text:
        recording = False
    if recording:
        print(f"P ({p.style.name}): {p.text}")

print("\n=== SECTION 5 IN MAU REPORT.docx ===")
recording = False
for p in doc.paragraphs:
    if "5." in p.text and "PHẦN MỞ RỘNG" in p.text:
        recording = True
    if recording:
        print(f"P ({p.style.name}): {p.text}")

print("\n=== ALL TABLES IN SECTION 5 ===")
# Tables near the end
for idx, table in enumerate(doc.tables[25:]):
    print(f"\n--- Table {idx+26} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for r_idx, row in enumerate(table.rows):
        row_str = " | ".join([cell.text.replace('\n', ' ') for cell in row.cells])
        print(f"Row {r_idx}: {row_str[:150]}")
