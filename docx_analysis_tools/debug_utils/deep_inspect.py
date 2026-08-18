#!/usr/bin/env python3
import docx
import json

doc = docx.Document('MAU REPORT.docx')

print("=== ALL HEADINGS ===")
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        print(f"[{p.style.name}] {p.text}")

print("\n=== SAMPLE VULNERABILITY PRESENTATION (Section 3.1 & 3.2) ===")
# Print paragraphs between Heading 3.1 and Heading 3.3
recording = False
for p in doc.paragraphs:
    if "3.1." in p.text:
        recording = True
    elif "3.3." in p.text:
        recording = False
    if recording:
        print(f"P ({p.style.name}): {p.text}")

print("\n=== SAMPLE TABLES IN MAU REPORT ===")
print(f"Total tables: {len(doc.tables)}")
for idx, table in enumerate(doc.tables[:5]):
    print(f"\n--- Table {idx+1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for r_idx, row in enumerate(table.rows):
        row_str = " | ".join([cell.text.replace('\n', ' ') for cell in row.cells])
        print(f"Row {r_idx}: {row_str[:150]}")
