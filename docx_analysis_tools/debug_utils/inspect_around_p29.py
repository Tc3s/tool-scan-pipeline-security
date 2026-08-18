#!/usr/bin/env python3
import docx

doc = docx.Document('MAU REPORT.docx')

for idx in range(15, 45):
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        print(f"P {idx} ({p.style.name}): {p.text}")
