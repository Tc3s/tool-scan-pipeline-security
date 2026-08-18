#!/usr/bin/env python3
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import re
import os
import json
import matplotlib.pyplot as plt

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def force_font(run, font_name="Times New Roman", size_pt=11, bold=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def enforce_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None:
        trPr.append(OxmlElement('w:cantSplit'))

def enforce_header_repeat(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:tblHeader')) is None:
        trPr.append(OxmlElement('w:tblHeader'))

def clean_html(text):
    if not isinstance(text, str) or pd.isna(text):
        return ""
    cleaned = re.sub(r'<[^>]+>', '', text)
    cleaned = cleaned.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
    return cleaned.strip()

SEV_MAP = {
    'CRITICAL': 'NGHIÊM TRỌNG',
    'HIGH': 'CAO',
    'MEDIUM': 'TRUNG BÌNH',
    'LOW': 'THẤP'
}

SEV_COLOR = {
    'NGHIÊM TRỌNG': RGBColor(255, 0, 0),
    'CAO': RGBColor(192, 0, 0),
    'TRUNG BÌNH': RGBColor(237, 125, 49),
    'THẤP': RGBColor(68, 114, 196)
}

# --- LOAD PERSISTENT KNOWLEDGE BASE ---
kb_path = 'docx_analysis_tools/full_vulnerability_knowledge_base.json'
with open(kb_path, 'r', encoding='utf-8') as f:
    db = json.load(f)

laptop_findings = db["laptop_findings"]
webserver_findings = db["webserver_findings"]

# BASELINE DATA (3 Months Ago Baseline for BOTH targets)
past_laptop = {'NGHIÊM TRỌNG': 1, 'CAO': 2, 'TRUNG BÌNH': 3, 'THẤP': 2}
past_webserver = {'NGHIÊM TRỌNG': 2, 'CAO': 4, 'TRUNG BÌNH': 6, 'THẤP': 3}

def get_counts_from_db(findings_list):
    counts = {'NGHIÊM TRỌNG': 0, 'CAO': 0, 'TRUNG BÌNH': 0, 'THẤP': 0}
    for item in findings_list:
        sev = item['severity_vi']
        if sev in counts:
            counts[sev] += 1
    return counts

curr_laptop = get_counts_from_db(laptop_findings)
curr_webserver = get_counts_from_db(webserver_findings)

# Configure Matplotlib fonts and styles
plt.rcParams['font.sans-serif'] = 'DejaVu Sans'
plt.rcParams['axes.edgecolor'] = '#cccccc'
plt.rcParams['axes.linewidth'] = 0.8

categories_vi = ['NGHIÊM TRỌNG', 'CAO', 'TRUNG BÌNH', 'THẤP']
colors_bar = ['#FF0000', '#C00000', '#ED7D31', '#4472C4']

# --- CHART 1: LAPTOP CHART (SECTION 2.1) ---
fig, ax = plt.subplots(figsize=(8.5, 4.5))
laptop_vals = [curr_laptop[c] for c in categories_vi]

bars1 = ax.bar(categories_vi, laptop_vals, color=colors_bar, width=0.45, edgecolor='#333333', linewidth=0.6)
ax.set_xlabel('Mức độ nghiêm trọng (Severity)', fontweight='bold', fontsize=11, labelpad=10)
ax.set_ylabel('Số lượng lỗ hổng phát hiện', fontweight='bold', fontsize=11, labelpad=10)
ax.set_title('BIỂU ĐỒ PHÂN BỔ LỖ HỔNG - MÁY TÍNH XÁCH TAY (192.168.95.135)', fontweight='bold', fontsize=12, pad=15)
ax.set_ylim(0, max(laptop_vals) + 2 if max(laptop_vals) > 0 else 5)
ax.grid(axis='y', linestyle='--', alpha=0.5)

for bar in bars1:
    yval = bar.get_height()
    if yval > 0:
        ax.text(bar.get_x() + bar.get_width()/2.0, yval + 0.15, str(int(yval)), ha='center', va='bottom', fontweight='bold', fontsize=10)

chart_laptop_path = 'docx_analysis_tools/chart_laptop.png'
plt.savefig(chart_laptop_path, dpi=300, bbox_inches='tight', pad_inches=0.3)
plt.close()


# --- CHART 2: WEB SERVER CHART (SECTION 2.2) ---
fig, ax = plt.subplots(figsize=(8.5, 4.5))
web_vals = [curr_webserver[c] for c in categories_vi]

bars2 = ax.bar(categories_vi, web_vals, color=colors_bar, width=0.45, edgecolor='#333333', linewidth=0.6)
ax.set_xlabel('Mức độ nghiêm trọng (Severity)', fontweight='bold', fontsize=11, labelpad=10)
ax.set_ylabel('Số lượng lỗ hổng phát hiện', fontweight='bold', fontsize=11, labelpad=10)
ax.set_title('BIỂU ĐỒ PHÂN BỔ LỖ HỔNG - WEB SERVER (192.168.95.138)', fontweight='bold', fontsize=12, pad=15)
ax.set_ylim(0, max(web_vals) + 3 if max(web_vals) > 0 else 5)
ax.grid(axis='y', linestyle='--', alpha=0.5)

for bar in bars2:
    yval = bar.get_height()
    if yval > 0:
        ax.text(bar.get_x() + bar.get_width()/2.0, yval + 0.35, str(int(yval)), ha='center', va='bottom', fontweight='bold', fontsize=10)

chart_web_path = 'docx_analysis_tools/chart_webserver.png'
plt.savefig(chart_web_path, dpi=300, bbox_inches='tight', pad_inches=0.3)
plt.close()


# --- BUILD DOCX ---
template_path = 'MAU REPORT.docx'
if os.path.exists(template_path):
    doc = docx.Document(template_path)
    # Clean body XML completely
    for child in list(doc._element.body):
        if not child.tag.endswith('sectPr'):
            doc._element.body.remove(child)
else:
    doc = docx.Document()
    section = doc.sections[0]
    section.page_width = docx.shared.Mm(210)
    section.page_height = docx.shared.Mm(297)
    section.top_margin = docx.shared.Mm(20)
    section.bottom_margin = docx.shared.Mm(20)
    section.left_margin = docx.shared.Mm(30)
    section.right_margin = docx.shared.Mm(20)

def add_h1(text):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    force_font(r, "Times New Roman", 16, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_h2(text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    force_font(r, "Times New Roman", 13, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_h3(text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    force_font(r, "Times New Roman", 12, bold=True, color_rgb=RGBColor(0, 0, 0))
    return p

def add_p(text, style='Normal', bold=False, color_rgb=None, size_pt=11):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    force_font(r, "Times New Roman", size_pt, bold=bold, color_rgb=color_rgb)
    return p

def format_table_cell(cell, text, bold=False, color_rgb=None, fill_hex=None, size_pt=10):
    if fill_hex:
        set_cell_background(cell, fill_hex)
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    force_font(r, "Times New Roman", size_pt, bold=bold, color_rgb=color_rgb)

# TITLE
p_title = doc.add_paragraph(style='Heading 1')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(24)
p_title.paragraph_format.space_after = Pt(18)
p_title.paragraph_format.keep_with_next = True
r_t = p_title.add_run("BÁO CÁO NỘI DUNG ĐÁNH GIÁ LỖ HỔNG BẢO MẬT")
force_font(r_t, "Times New Roman", 18, bold=True, color_rgb=RGBColor(0, 32, 96))

# --- TABLE OF CONTENTS SECTION ---
add_h1("MỤC LỤC")
t_toc = doc.add_table(rows=1, cols=2)
t_toc.style = 'Table Grid'
t_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
enforce_row_cant_split(t_toc.rows[0])
enforce_header_repeat(t_toc.rows[0])
format_table_cell(t_toc.rows[0].cells[0], "NỘI DUNG CẤU TRÚC BÁO CÁO", bold=True, fill_hex='D9D9D9')
format_table_cell(t_toc.rows[0].cells[1], "TRANG", bold=True, fill_hex='D9D9D9')

toc_items = [
    ("1. THÔNG TIN DỰ ÁN", "3"),
    ("   1.1. PHIÊN BẢN TÀI LIỆU", "3"),
    ("   1.2. THỜI GIAN TRIỂN KHAI", "3"),
    ("   1.3. NHÂN SỰ TRIỂN KHAI", "3"),
    ("   1.4. MỤC ĐÍCH ĐÁNH GIÁ", "4"),
    ("   1.5. NỘI DUNG ĐÁNH GIÁ", "4"),
    ("   1.6. PHẠM VI THỰC HIỆN", "4"),
    ("2. BÁO CÁO TỔNG QUÁT", "5"),
    ("   2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", "5"),
    ("   2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER (192.168.95.138)", "6"),
    ("   2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA / TRENDING)", "8"),
    ("   2.4. KHUYẾN NGHỊ TỔNG QUAN", "9"),
    ("3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)", "10"),
    ("4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER (192.168.95.138)", "12"),
    ("5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO", "25"),
    ("   5.1. MÔ HÌNH ĐÁNH GIÁ RỦI RO THEO OWASP RISK RATING", "25"),
    ("   5.2. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0", "26"),
    ("   5.3. THANG ĐIỂM ĐÁNH GIÁ CVSS V3.0", "27"),
]

for section_name, page_num in toc_items:
    row = t_toc.add_row()
    enforce_row_cant_split(row)
    r_c = row.cells
    is_main = not section_name.startswith("   ")
    format_table_cell(r_c[0], section_name, bold=is_main)
    format_table_cell(r_c[1], page_num, bold=is_main)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Native Word TOC Field
p_native_toc = doc.add_paragraph()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "TOC \\o \"1-3\" \\h \\z \\u"
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
r1 = OxmlElement('w:r'); r1.append(fldChar1); p_native_toc._element.append(r1)
r2 = OxmlElement('w:r'); r2.append(instrText); p_native_toc._element.append(r2)
r3 = OxmlElement('w:r'); r3.append(fldChar2); p_native_toc._element.append(r3)
r4 = OxmlElement('w:r'); r4.append(fldChar3); p_native_toc._element.append(r4)

# 1. THÔNG TIN DỰ ÁN
add_h1("1. THÔNG TIN DỰ ÁN")

add_h2("1.1. PHIÊN BẢN TÀI LIỆU")
t_ver = doc.add_table(rows=4, cols=2)
t_ver.style = 'Table Grid'
t_ver.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in t_ver.rows: enforce_row_cant_split(row)
enforce_header_repeat(t_ver.rows[0])

ver_data = [
    ("Phiên bản tài liệu", "1.0"),
    ("Ngày báo cáo", "10/08/2026"),
    ("Khách hàng", "[Tên Khách Hàng]"),
    ("Công ty cung ứng dịch vụ", "CÔNG TY CỔ PHẦN CÔNG NGHỆ DTG")
]
for idx, (lbl, val) in enumerate(ver_data):
    format_table_cell(t_ver.rows[idx].cells[0], lbl, bold=True, fill_hex='F2F2F2')
    format_table_cell(t_ver.rows[idx].cells[1], val)

add_h2("1.2. THỜI GIAN TRIỂN KHAI")
add_p("Thời gian thực hiện đánh giá: 01/08/2026 – 10/08/2026")

add_h2("1.3. NHÂN SỰ TRIỂN KHAI")
t_staff = doc.add_table(rows=2, cols=3)
t_staff.style = 'Table Grid'
t_staff.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in t_staff.rows: enforce_row_cant_split(row)
enforce_header_repeat(t_staff.rows[0])

format_table_cell(t_staff.rows[0].cells[0], "STT", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[0].cells[1], "Nhân sự", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[0].cells[2], "Vai Trò", bold=True, fill_hex='D9D9D9')
format_table_cell(t_staff.rows[1].cells[0], "1")
format_table_cell(t_staff.rows[1].cells[1], "Chuyên gia An toàn thông tin")
format_table_cell(t_staff.rows[1].cells[2], "Kiểm thử viên chính (Lead Pentester)")

add_h2("1.4. MỤC ĐÍCH ĐÁNH GIÁ")
add_p("• Kiểm tra và đánh giá toàn diện hiện trạng an toàn thông tin đối với các thiết bị máy tính xách tay và hệ thống máy chủ dịch vụ Web Server định kỳ 3 tháng/lần.")
add_p("• Kịp thời phát hiện, phân loại và cảnh báo các lỗ hổng bảo mật, điểm yếu cấu hình hoặc các phần mềm lỗi thời tồn tại trên các mục tiêu.")
add_p("• Đề xuất phương án khắc phục, giảm thiểu rủi ro an ninh mạng làm cơ sở tham mưu cho đơn vị.")

add_h2("1.5. NỘI DUNG ĐÁNH GIÁ")
add_p("• Rà quét tự động kết hợp phân tích thủ công để xác định các cổng/dịch vụ mạng đang mở, điểm yếu giao thức và lỗ hổng phần mềm.")
add_p("• Phân loại mức độ nghiêm trọng của lỗ hổng dựa trên tiêu chuẩn quốc tế CVSS v3.0 và mô hình OWASP Risk Rating.")
add_p("• Lập báo cáo kết quả chi tiết, tư vấn phương án khắc phục và tái kiểm tra.")

add_h2("1.6. PHẠM VI THỰC HIỆN")
add_p("Hoạt động rà quét được thực hiện đối với 2 mục tiêu chính thông qua phương thức kiểm thử từ bên ngoài (External Scan - Internet Simulation):")

t_scope = doc.add_table(rows=3, cols=4)
t_scope.style = 'Table Grid'
t_scope.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in t_scope.rows: enforce_row_cant_split(row)
enforce_header_repeat(t_scope.rows[0])

scope_headers = ['STT', 'Thông tin thiết bị', 'Địa chỉ IP', 'Phương thức đánh giá']
for i, h in enumerate(scope_headers):
    format_table_cell(t_scope.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

format_table_cell(t_scope.rows[1].cells[0], "1")
format_table_cell(t_scope.rows[1].cells[1], "Máy tính xách tay Windows 11")
format_table_cell(t_scope.rows[1].cells[2], "192.168.95.135")
format_table_cell(t_scope.rows[1].cells[3], "Rà quét từ bên ngoài (Internet Simulation)")

format_table_cell(t_scope.rows[2].cells[0], "2")
format_table_cell(t_scope.rows[2].cells[1], "Hệ thống Web Server")
format_table_cell(t_scope.rows[2].cells[2], "192.168.95.138")
format_table_cell(t_scope.rows[2].cells[3], "Rà quét từ bên ngoài (Internet Simulation)")


# 2. BÁO CÁO TỔNG QUÁT
add_h1("2. BÁO CÁO TỔNG QUÁT")

# --- SECTION 2.1: LAPTOP WITH CHART 1 ---
add_h2("2.1. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)")
add_p(f"Tổng số lỗ hổng phát hiện trên Máy tính xách tay Windows 11: {len(laptop_findings)} lỗ hổng.")

doc.add_picture(chart_laptop_path, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

t_lap = doc.add_table(rows=1, cols=4)
t_lap.style = 'Table Grid'
t_lap.alignment = WD_TABLE_ALIGNMENT.CENTER
enforce_row_cant_split(t_lap.rows[0])
enforce_header_repeat(t_lap.rows[0])

for i, h in enumerate(['STT', 'Mức độ', 'Lỗ hổng', 'Thiết bị']):
    format_table_cell(t_lap.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

for item in laptop_findings:
    row = t_lap.add_row()
    enforce_row_cant_split(row)
    r_cells = row.cells
    format_table_cell(r_cells[0], str(item['id']))
    sev_vi = item['severity_vi']
    color = SEV_COLOR.get(sev_vi, RGBColor(0,0,0))
    format_table_cell(r_cells[1], sev_vi, bold=True, color_rgb=color)
    format_table_cell(r_cells[2], item['finding_name'])
    format_table_cell(r_cells[3], "Máy tính xách tay Windows 11")


# --- SECTION 2.2: WEB SERVER WITH CHART 2 ---
add_h2("2.2. KẾT QUẢ DÒ QUÉT LỖ HỔNG BẢO MẬT WEB SERVER (192.168.95.138)")
add_p(f"Tổng số lỗ hổng phát hiện trên Web Server: {len(webserver_findings)} lỗ hổng.")

doc.add_picture(chart_web_path, width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

t_web = doc.add_table(rows=1, cols=4)
t_web.style = 'Table Grid'
t_web.alignment = WD_TABLE_ALIGNMENT.CENTER
enforce_row_cant_split(t_web.rows[0])
enforce_header_repeat(t_web.rows[0])

for i, h in enumerate(['STT', 'Mức độ', 'Lỗ hổng', 'Thiết bị']):
    format_table_cell(t_web.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

for item in webserver_findings:
    row = t_web.add_row()
    enforce_row_cant_split(row)
    r_cells = row.cells
    format_table_cell(r_cells[0], str(item['id']))
    sev_vi = item['severity_vi']
    color = SEV_COLOR.get(sev_vi, RGBColor(0,0,0))
    format_table_cell(r_cells[1], sev_vi, bold=True, color_rgb=color)
    format_table_cell(r_cells[2], item['finding_name'])
    format_table_cell(r_cells[3], "Web Server")


# --- SECTION 2.3: DELTA & TRENDING ---
add_h2("2.3. SO SÁNH VỚI CHU KỲ 3 THÁNG TRƯỚC (DELTA / TRENDING)")
add_p("Bảng chi tiết thống kê biến động Delta lỗ hổng theo từng mục tiêu giữa Chu kỳ 3 tháng trước và Chu kỳ hiện tại:")

t_delta = doc.add_table(rows=1, cols=6)
t_delta.style = 'Table Grid'
t_delta.alignment = WD_TABLE_ALIGNMENT.CENTER
enforce_row_cant_split(t_delta.rows[0])
enforce_header_repeat(t_delta.rows[0])

delta_headers = ['Chu kỳ rà quét', 'Mục tiêu', 'Nghiêm trọng', 'Cao', 'Trung bình', 'Thấp']
for i, h in enumerate(delta_headers):
    format_table_cell(t_delta.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

delta_rows = [
    ("Chu kỳ 3 tháng trước", "Laptop (192.168.95.135)", past_laptop),
    ("Chu kỳ hiện tại", "Laptop (192.168.95.135)", curr_laptop),
    ("Chu kỳ 3 tháng trước", "Web Server (192.168.95.138)", past_webserver),
    ("Chu kỳ hiện tại", "Web Server (192.168.95.138)", curr_webserver),
]

for cycle, tgt, data_dict in delta_rows:
    row = t_delta.add_row()
    enforce_row_cant_split(row)
    r_cells = row.cells
    format_table_cell(r_cells[0], cycle)
    format_table_cell(r_cells[1], tgt)
    format_table_cell(r_cells[2], str(data_dict['NGHIÊM TRỌNG']), bold=(data_dict['NGHIÊM TRỌNG'] > 0), color_rgb=RGBColor(255,0,0) if data_dict['NGHIÊM TRỌNG'] > 0 else None)
    format_table_cell(r_cells[3], str(data_dict['CAO']), bold=(data_dict['CAO'] > 0), color_rgb=RGBColor(192,0,0) if data_dict['CAO'] > 0 else None)
    format_table_cell(r_cells[4], str(data_dict['TRUNG BÌNH']), bold=(data_dict['TRUNG BÌNH'] > 0), color_rgb=RGBColor(237,125,49) if data_dict['TRUNG BÌNH'] > 0 else None)
    format_table_cell(r_cells[5], str(data_dict['THẤP']))

add_p("Phân tích xu hướng rủi ro (Trending Analysis):", bold=True)
add_p("• Đối với Máy tính xách tay Windows 11 (192.168.95.135): Số lượng lỗ hổng giảm từ 8 xuống 4 lỗi, đặc biệt các lỗi Nghiêm trọng đã được triệt tiêu hoàn toàn. Cho thấy công tác cập nhật bản vá hệ điều hành đã có hiệu quả tích cực.")
add_p("• Đối với Hệ thống Web Server (192.168.95.138): Mức độ rủi ro tăng mạnh. Số lỗi Nghiêm trọng tăng từ 2 lên 10 lỗi, lỗi Cao giữ ở mức 4 lỗi. Sự gia tăng này cảnh báo bề mặt tấn công công cộng đang bị đe dọa trực tiếp bởi các lỗ hổng mã nguồn ứng dụng và dịch vụ chưa được vá.")

add_h2("2.4. KHUYẾN NGHỊ TỔNG QUAN")
add_p("1. Ưu tiên xử lý khẩn cấp 10 lỗ hổng NGHIÊM TRỌNG và 4 lỗ hổng CAO trên Web Server trong vòng 24 - 48 giờ.")
add_p("2. Thực hiện nâng cấp các phiên bản dịch vụ Web Server và thư viện liên quan lên phiên bản an toàn mới nhất.")
add_p("3. Thiết lập hệ thống Tường lửa ứng dụng Web (WAF) để ngăn chặn các cuộc tấn công khai thác lỗ hổng từ môi trường Internet.")
add_p("4. Duy trì chính sách rà quét lỗ hổng bảo mật định kỳ 3 tháng/lần để phát hiện sớm các rủi ro phát sinh.")


# HELPER FOR VULNERABILITY DETAILS
def render_detailed_vuln_item(sec_num, item):
    add_h3(f"{sec_num}.{item['id']}. {item['finding_name'].upper()}")
    
    add_p(f"Mô tả: {item['clean_description_vi']}")
    
    p_sev = doc.add_paragraph(style='Normal')
    p_sev.paragraph_format.space_after = Pt(3)
    r_lbl = p_sev.add_run("Mức độ: ")
    force_font(r_lbl, "Times New Roman", 11, bold=False)
    
    sev_vi = item['severity_vi']
    color = SEV_COLOR.get(sev_vi, RGBColor(0,0,0))
    r_sev = p_sev.add_run(sev_vi)
    force_font(r_sev, "Times New Roman", 11, bold=True, color_rgb=color)

    add_p(f"CVSS: {item['cvss']}")
    add_p(f"CVE: {item['cve']}")
    
    add_p("Bằng chứng:", bold=True)
    t_ev = doc.add_table(rows=1, cols=1)
    t_ev.style = 'Table Grid'
    t_ev.alignment = WD_TABLE_ALIGNMENT.CENTER
    enforce_row_cant_split(t_ev.rows[0])
    format_table_cell(t_ev.rows[0].cells[0], item['clean_evidence_vi'], fill_hex='F9F9F9', size_pt=10)
    
    add_p(f"Khuyến nghị: {item['clean_solution_vi']}")


# 3. BÁO CÁO CHI TIẾT LAPTOP
add_h1("3. BÁO CÁO CHI TIẾT DÒ QUÉT LỖ HỔNG BẢO MẬT MÁY TÍNH XÁCH TAY (192.168.95.135)")
for item in laptop_findings:
    render_detailed_vuln_item(3, item)

# 4. BÁO CÁO CHI TIẾT WEB SERVER
add_h1("4. BÁO CÁO ĐÁNH GIÁ CHI TIẾT WEB SERVER (192.168.95.138)")
for item in webserver_findings:
    render_detailed_vuln_item(4, item)


# 5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO
add_h1("5. PHẦN MỞ RỘNG: PHÂN LOẠI RỦI RO")

add_h2("5.1. MÔ HÌNH ĐÁNH GIÁ RỦI RO THEO OWASP RISK RATING")
add_p("Mô hình đánh giá rủi ro an toàn thông tin dựa trên tiêu chuẩn OWASP Risk Rating Methodology, kết hợp 2 yếu tố chính: Khả năng xảy ra (Likelihood) và Mức độ ảnh hưởng (Impact).")

# Table 31: OWASP Matrix
t_owasp = doc.add_table(rows=6, cols=5)
t_owasp.style = 'Table Grid'
t_owasp.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in t_owasp.rows: enforce_row_cant_split(row)
enforce_header_repeat(t_owasp.rows[0])

owasp_headers = ["MA TRẬN OWASP", "KHÔNG", "THẤP", "TRUNG BÌNH", "CAO"]
for i, h in enumerate(owasp_headers):
    format_table_cell(t_owasp.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

owasp_data = [
    ["MỨC ĐỘ ẢNH HƯỞNG", "CAO", "TRUNG BÌNH", "CAO", "NGHIÊM TRỌNG"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "TRUNG BÌNH", "THẤP", "TRUNG BÌNH", "CAO"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "THẤP", "KHÔNG", "THẤP", "TRUNG BÌNH"],
    ["MỨC ĐỘ ẢNH HƯỞNG", "-", "THẤP", "TRUNG BÌNH", "CAO"],
    ["-", "KHẢ NĂNG XẢY RA", "THẤP", "TRUNG BÌNH", "CAO"]
]

for r_idx, row in enumerate(owasp_data):
    for c_idx, val in enumerate(row):
        cell = t_owasp.rows[r_idx+1].cells[c_idx]
        color = SEV_COLOR.get(val, None)
        format_table_cell(cell, val, bold=(color is not None), color_rgb=color)

add_h2("5.2. MÔ HÌNH ĐÁNH GIÁ MỨC ĐỘ NGHIÊM TRỌNG THEO CVSS V3.0")
add_p("Common Vulnerability Scoring System (CVSS) v3.0 là tiêu chuẩn quốc tế mở dùng để đánh giá mức độ nghiêm trọng của lỗ hổng phần mềm.")

add_h3("5.2.1. EXPLOITABILITY METRICS")
add_p("• Attack Vector (AV): Đánh giá môi trường khai thác lỗ hổng (Network, Adjacent, Local, Physical).")
add_p("• Attack Complexity (AC): Đánh giá độ phức tạp của cuộc tấn công (Low, High).")
add_p("• Privileges Required (PR): Đánh giá mức độ quyền hạn cần thiết của kẻ tấn công (None, Low, High).")
add_p("• User Interaction (UI): Đánh giá yêu cầu sự tương tác của người dùng (None, Required).")

add_h3("5.2.2. IMPACT METRICS")
add_p("• Confidentiality Impact (C): Đánh giá mức độ rò rỉ thông tin bảo mật.")
add_p("• Integrity Impact (I): Đánh giá mức độ sai lệch hoặc sửa đổi dữ liệu.")
add_p("• Availability Impact (A): Đánh giá mức độ gián đoạn khả năng cung cấp dịch vụ.")

add_h2("5.3. THANG ĐIỂM ĐÁNH GIÁ CVSS V3.0")
t_cvss = doc.add_table(rows=5, cols=3)
t_cvss.style = 'Table Grid'
t_cvss.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in t_cvss.rows: enforce_row_cant_split(row)
enforce_header_repeat(t_cvss.rows[0])

cvss_headers = ["Điểm CVSS v3.0", "Mức độ nghiêm trọng", "Mô tả rủi ro"]
for i, h in enumerate(cvss_headers):
    format_table_cell(t_cvss.rows[0].cells[i], h, bold=True, fill_hex='D9D9D9')

cvss_table_data = [
    ("9.0 – 10.0", "NGHIÊM TRỌNG (Critical)", "Lỗ hổng cho phép khai thác từ xa không cần xác thực, chiếm quyền kiểm soát hệ thống."),
    ("7.0 – 8.9", "CAO (High)", "Lỗ hổng gây ảnh hưởng nghiêm trọng đến tính bảo mật hoặc sẵn sàng của hệ thống."),
    ("4.0 – 6.9", "TRUNG BÌNH (Medium)", "Lỗ hổng yêu cầu điều kiện khai thác phức tạp hoặc quyền truy cập hạn chế."),
    ("0.1 – 3.9", "THẤP (Low)", "Lỗ hổng có mức độ tác động nhỏ, khó khai thác thành công.")
]

for idx, (score, sev, desc) in enumerate(cvss_table_data):
    r_cells = t_cvss.rows[idx+1].cells
    format_table_cell(r_cells[0], score)
    
    sev_key = sev.split()[0]
    color = SEV_COLOR.get(sev_key, None)
    
    format_table_cell(r_cells[1], sev, bold=True, color_rgb=color)
    format_table_cell(r_cells[2], desc)

doc.save('Bao_Cao_An_Toan_Thong_Tin_2026.docx')
print("Successfully generated PERFECTION Bao_Cao_An_Toan_Thong_Tin_2026.docx")
